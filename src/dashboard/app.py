"""
Дашборд проекта neuropulse-grant-agent.
Блоки: диалог с агентом, ближайшие сроки, аудит чата, статус агента, быстрые действия, Vector Store, ссылки.
При нажатии на быстрые действия формулировка уходит в диалог с агентом.
"""

from __future__ import annotations

import hashlib
from typing import Any
import sys
from pathlib import Path
from datetime import datetime, timedelta

_PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
if str(_PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(_PROJECT_ROOT))

import json
import logging
import os

import streamlit as st

from src.agent_api_client import load_config, PROJECT_ROOT

# Чтобы дашборд видел YANDEX_VECTOR_STORE_ID из config/.env
try:
    from dotenv import load_dotenv
    load_dotenv(PROJECT_ROOT / "config" / ".env")
    load_dotenv(PROJECT_ROOT / ".env")
except ImportError:
    pass
from src.git_auto_push import git_push_if_changes, _register_atexit

logger = logging.getLogger(__name__)

AUDIT_LOG = PROJECT_ROOT / "data" / "chat_audit.jsonl"
CALENDAR_PATH = PROJECT_ROOT / "data" / "grant_calendar.json"
CALENDAR_EXAMPLE = PROJECT_ROOT / "data" / "grant_calendar.example.json"
CONTACTS_PATH = PROJECT_ROOT / "data" / "grant_contacts.json"
CONTACTS_EXAMPLE = PROJECT_ROOT / "data" / "grant_contacts.example.json"
GRANT_DASHBOARD_PATH = PROJECT_ROOT / "data" / "grant_project_dashboard.json"
GRANT_DASHBOARD_EXAMPLE = PROJECT_ROOT / "data" / "grant_project_dashboard.example.json"
KKT_PATH = PROJECT_ROOT / "data" / "grant_kkt.json"
KKT_EXAMPLE = PROJECT_ROOT / "data" / "grant_kkt.example.json"
RISKS_PATH = PROJECT_ROOT / "data" / "grant_risks.json"
RISKS_EXAMPLE = PROJECT_ROOT / "data" / "grant_risks.example.json"
COMMUNICATIONS_PATH = PROJECT_ROOT / "data" / "grant_communications.json"
COMMUNICATIONS_EXAMPLE = PROJECT_ROOT / "data" / "grant_communications.example.json"
DOCUMENTS_ARCHIVE_PATH = PROJECT_ROOT / "data" / "grant_documents_archive.json"
DOCUMENTS_ARCHIVE_EXAMPLE = PROJECT_ROOT / "data" / "grant_documents_archive.example.json"
NOTIFICATION_SETTINGS_PATH = PROJECT_ROOT / "data" / "notification_settings.json"
NOTIFICATION_SETTINGS_EXAMPLE = PROJECT_ROOT / "data" / "notification_settings.example.json"
DASHBOARD_LAYOUT_PATH = PROJECT_ROOT / "data" / "dashboard_layout.json"
MANUAL_CALENDAR_EVENTS_PATH = PROJECT_ROOT / "data" / "neuropulse_calendar_manual.json"

# Папка «Паспорт»: паспорт проекта и краткая история диалогов (автосохранение раз в 15 мин)
PASSPORT_DIR = PROJECT_ROOT / "Паспорт"
PASSPORT_PROJECT_FILE = PASSPORT_DIR / "паспорт_проекта.md"
PASSPORT_HISTORY_FILE = PASSPORT_DIR / "история_диалогов.md"
PASSPORT_LAST_SAVE_FILE = PASSPORT_DIR / ".last_passport_save"
PASSPORT_INTERVAL_MINUTES = 15

STATUS_LABELS = {"not_started": "⚪ Не начато", "in_progress": "🟡 В работе", "done": "🟢 Выполнено", "overdue": "🔴 Просрочено"}
STATUS_COLORS = {"not_started": "#888", "in_progress": "#d4a017", "done": "#28a745", "overdue": "#dc3545"}

# Постельные оттенки карточек: (фон карточки, фон заголовка — чуть ярче)
BLOCK_PALETTE: dict[str, tuple[str, str]] = {
    "grant_header":        ("#f0eeff", "#e2d9ff"),  # лаванда
    "grant_stages":        ("#e8f6ef", "#ceead9"),  # мята
    "kkt":                 ("#fff2e8", "#ffe0c6"),  # персик
    "grant_budget":        ("#fefae6", "#fdf3c8"),  # масло
    "grant_indicators":    ("#e6f3ff", "#cce5ff"),  # небесно-голубой
    "risks":               ("#fceaea", "#f9d2d2"),  # пудровый розовый
    "grant_info_activity": ("#e6faf6", "#cbf0e8"),  # тиффани
    "schedule":            ("#f0e8ff", "#e0d0ff"),  # сиреневый
    "calendar_month":      ("#e8f7ff", "#c9e9ff"),  # голубой
    "reminders":           ("#fff5e6", "#ffe9c8"),  # янтарный
    "neuropulse_cal":      ("#e6f9ff", "#c8efff"),  # циан
    "audit":               ("#eef5e8", "#daecd0"),  # шалфей
    "agent_stats":         ("#fce8f5", "#f8d0ec"),  # орхидея
    "status":              ("#eaf1f8", "#d2e4f2"),  # стальной
    "tools":               ("#fef6e8", "#faead0"),  # кремовый
    "contacts":            ("#ede8fa", "#ddd0f6"),  # пыльно-сиреневый
    "notifications":       ("#fff8e6", "#ffedc8"),  # медовый
    "communications":      ("#e8f8f0", "#c8f0da"),  # морской
    "documents_archive":   ("#e8f2fa", "#cce0f4"),  # васильковый
    "links":               ("#fce8ed", "#f8d0da"),  # коралловый
    "vector_store":        ("#e8faf6", "#c8f2e8"),  # аквамарин
    "parser":              ("#ece8fa", "#dcd0f4"),  # индиго
}


def _inject_block_palette_css() -> None:
    """Генерирует и инжектирует CSS-правила постельных оттенков для каждого блока."""
    rules: list[str] = []
    for bid, (card_bg, hdr_bg) in BLOCK_PALETTE.items():
        sel = f'[data-testid="stVerticalBlock"]:has(.np-bm-{bid})'
        rules.append(
            f'{sel} {{'
            f' background: linear-gradient(to bottom, {hdr_bg} 0px, {hdr_bg} 62px, {card_bg} 62px, {card_bg} 100%) !important;'
            f' border-color: {hdr_bg} !important;'
            f'}}'
        )
    # Выравнивание заголовков и маркеры — в neuropulse_theme.css; здесь только палитра блоков
    css = "\n".join(rules)
    st.markdown(f'<style id="np-block-palette">\n{css}\n</style>', unsafe_allow_html=True)


def _block_header(title: str, block_id: str) -> bool:
    """Заголовок блока: название и шестерёнка (настройки). Перемещение блоков — только в экспандере «Расположение блоков»."""
    key_open = f"dashboard_settings_{block_id}"
    if key_open not in st.session_state:
        st.session_state[key_open] = False

    col1, col2 = st.columns([5, 1])
    with col1:
        st.subheader(title)
    with col2:
        if st.button("⚙️", key=f"gear_{block_id}", help="Настройки блока"):
            st.session_state[key_open] = not st.session_state[key_open]
            st.rerun()
    return st.session_state[key_open]


def _block_with_settings(title: str, block_id: str, render_content, render_settings):
    """Рендер блока: заголовок + шестерёнка, при открытых настройках — экспандер, затем контент."""
    with st.container(border=True):
        # Невидимый маркер для CSS :has() — позволяет красить карточку по block_id
        st.markdown(f'<div class="np-bm-{block_id}" style="display:none;height:0;overflow:hidden"></div>', unsafe_allow_html=True)
        settings_open = _block_header(title, block_id)
        if settings_open:
            with st.expander("Настройки блока", expanded=True):
                render_settings(block_id)
                st.divider()
                is_hidden = _is_block_hidden(block_id)
                hide_label = "👁 Показать блок на дашборде" if is_hidden else "🙈 Скрыть блок с дашборда"
                if st.button(hide_label, key=f"toggle_hide_{block_id}"):
                    hidden = _load_hidden_blocks()
                    if is_hidden:
                        hidden.discard(block_id)
                    else:
                        hidden.add(block_id)
                    _save_hidden_blocks(hidden)
                    st.session_state["dashboard_hidden_blocks"] = hidden
                    st.session_state[f"dashboard_settings_{block_id}"] = False
                    st.rerun()
                if st.button("Закрыть настройки", key=f"close_{block_id}"):
                    st.session_state[f"dashboard_settings_{block_id}"] = False
                    st.rerun()
        render_content(block_id)


def _apply_block_move(
    action: str, bid: str, side: str, index: int, total: int,
    left_ids: list[str], right_ids: list[str],
) -> None:
    """Перемещает блок: action in ('left','right','up','down')."""
    if action == "left" and side == "right" and bid in right_ids:
        right_ids.remove(bid)
        left_ids.append(bid)
    elif action == "right" and side == "left" and bid in left_ids:
        left_ids.remove(bid)
        right_ids.append(bid)
    elif action == "up" and index > 0:
        if side == "left":
            left_ids[index], left_ids[index - 1] = left_ids[index - 1], left_ids[index]
        else:
            right_ids[index], right_ids[index - 1] = right_ids[index - 1], right_ids[index]
    elif action == "down" and index < total - 1:
        if side == "left":
            left_ids[index], left_ids[index + 1] = left_ids[index + 1], left_ids[index]
        else:
            right_ids[index], right_ids[index + 1] = right_ids[index + 1], right_ids[index]


def _render_move_arrows(bid: str, side: str, index: int, total: int, left_ids: list[str], right_ids: list[str]) -> None:
    """Рендерит строку с кнопками ← ↑ ↓ → для переноса блока."""
    c0, c1, c2, c3 = st.columns(4)
    with c0:
        if st.button("←", key=f"arr_l_{bid}", help="В левую колонку", disabled=(side != "right")):
            _apply_block_move("left", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()
    with c1:
        if st.button("↑", key=f"arr_u_{bid}", help="Выше", disabled=(index <= 0)):
            _apply_block_move("up", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()
    with c2:
        if st.button("↓", key=f"arr_d_{bid}", help="Ниже", disabled=(index >= total - 1)):
            _apply_block_move("down", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()
    with c3:
        if st.button("→", key=f"arr_r_{bid}", help="В правую колонку", disabled=(side != "left")):
            _apply_block_move("right", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()


def _block_header_with_arrows(
    title: str, bid: str, side: str, index: int, total: int,
    left_ids: list[str], right_ids: list[str],
) -> None:
    """Заголовок блока: название и стрелки ← ↑ ↓ → для переноса."""
    col_title, c0, c1, c2, c3 = st.columns([3, 1, 1, 1, 1])
    with col_title:
        st.subheader(title)
    with c0:
        if st.button("←", key=f"arr_l_{bid}", help="В левую колонку", disabled=(side != "right")):
            _apply_block_move("left", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()
    with c1:
        if st.button("↑", key=f"arr_u_{bid}", help="Выше", disabled=(index <= 0)):
            _apply_block_move("up", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()
    with c2:
        if st.button("↓", key=f"arr_d_{bid}", help="Ниже", disabled=(index >= total - 1)):
            _apply_block_move("down", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()
    with c3:
        if st.button("→", key=f"arr_r_{bid}", help="В правую колонку", disabled=(side != "left")):
            _apply_block_move("right", bid, side, index, total, left_ids, right_ids)
            _save_dashboard_layout(left_ids, right_ids)
            st.rerun()


def _get(key: str, default):
    if key not in st.session_state:
        st.session_state[key] = default
    return st.session_state[key]


def _load_grant_dashboard_data() -> dict:
    """Загружает данные мониторинга гранта из JSON."""
    for path in (GRANT_DASHBOARD_PATH, GRANT_DASHBOARD_EXAMPLE):
        if path.exists():
            try:
                with open(path, encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                logger.warning("Ошибка чтения %s: %s", path.name, e)
    return {}


# ---------- Диалог с агентом (сессия и аудит) ----------
def _ensure_conversation() -> str | None:
    """Возвращает conversation_id для диалога с агентом; при необходимости создаёт или загружает из файла. None если API не настроен."""
    try:
        from dotenv import load_dotenv
        load_dotenv(PROJECT_ROOT / "config" / ".env")
        load_dotenv(PROJECT_ROOT / ".env")
    except ImportError:
        pass
    try:
        from src.yandex_ai_client import load_session_id, create_conversation, save_session_id
        conv_id = load_session_id()
        if conv_id:
            return conv_id
        new_id = create_conversation()
        save_session_id(new_id)
        return new_id
    except Exception as e:
        logger.warning("Агент не настроен для диалога: %s", e)
        return None


def _log_audit(conversation_id: str, prompt_len: int, reply_len: int) -> None:
    """Запись в data/chat_audit.jsonl для аудита."""
    try:
        AUDIT_LOG.parent.mkdir(parents=True, exist_ok=True)
        conv_hash = hashlib.sha256(conversation_id.encode()).hexdigest()[:16]
        record = {
            "ts": datetime.utcnow().isoformat() + "Z",
            "conv_id_hash": conv_hash,
            "prompt_len": prompt_len,
            "reply_len": reply_len,
        }
        with open(AUDIT_LOG, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")
    except Exception as e:
        logger.warning("Ошибка записи аудит-лога: %s", e)


def _agent_config_hint() -> str:
    """Краткая подсказка, что не задано (без вывода секретов)."""
    try:
        from dotenv import load_dotenv
        load_dotenv(PROJECT_ROOT / "config" / ".env")
        load_dotenv(PROJECT_ROOT / ".env")
    except ImportError:
        pass
    env_path = PROJECT_ROOT / "config" / ".env"
    parts = ["Проверьте **config/.env** (или корневой .env)."]
    if not env_path.exists():
        parts.append(f"Файл `config/.env` не найден по пути `{env_path}`.")
    key_ok = bool(os.environ.get("YANDEX_API_KEY", "").strip())
    folder_ok = bool(os.environ.get("YANDEX_FOLDER_ID", "").strip())
    if not key_ok:
        parts.append("**YANDEX_API_KEY** не задан или пуст.")
    if not folder_ok:
        parts.append("**YANDEX_FOLDER_ID** не задан или пуст (нужен идентификатор каталога, обычно b1g...).")
    if key_ok and folder_ok:
        parts.append("Ключ и каталог заданы — возможна ошибка сети или прав API (см. логи дашборда).")
    return " ".join(parts)


def _send_to_agent(prompt: str) -> str | None:
    """Отправляет сообщение агенту, возвращает ответ или строку с описанием ошибки."""
    conv_id = _ensure_conversation()
    if not conv_id:
        return "Не удалось подключиться к агенту. " + _agent_config_hint()
    try:
        from src.yandex_ai_client import ask, ask_in_conversation
        reply = ask_in_conversation(conv_id, prompt)
        if reply and "API вернул пустой ответ" in reply:
            try:
                fallback = ask(prompt)
                if fallback and fallback.strip():
                    return fallback + "\n\n_(Ответ получен одним запросом без сессии — при проблемах с диалогом.)_"
            except Exception as e:
                logger.debug("Fallback ask() не удался: %s", e)
        _log_audit(conv_id, len(prompt), len(reply or ""))
        return reply
    except Exception as e:
        logger.exception("Ошибка запроса к агенту: %s", e)
        return f"**Ошибка запроса:** {e}"


# ---------- Блок 1: Ближайшие сроки по гранту ----------
def _load_calendar_local(calendar_path: Path | None = None) -> list[dict]:
    path = calendar_path or _get("dashboard_schedule_calendar_path", "")
    if path:
        p = Path(path)
        if p.exists():
            try:
                with open(p, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
    if CALENDAR_PATH.exists():
        with open(CALENDAR_PATH, encoding="utf-8") as f:
            return json.load(f)
    if CALENDAR_EXAMPLE.exists():
        with open(CALENDAR_EXAMPLE, encoding="utf-8") as f:
            return json.load(f)
    return []


def _load_schedule_events() -> list[dict]:
    """Загружает события: из Яндекс Календаря (CalDAV) или единый набор (грант + ККТ + этапы + импорт из календаря) — по настройке блока."""
    source = _get("dashboard_schedule_source", "local")
    days = _get("dashboard_schedule_days", 30)
    min_year = _get("dashboard_schedule_min_year", 2026)
    today = datetime.utcnow().date()
    end = today + timedelta(days=days)

    if source == "yandex":
        try:
            from src.yandex_calendar_client import fetch_events, is_configured
            if not is_configured():
                return []
            return fetch_events(today, end)
        except Exception as e:
            logger.warning("Ошибка загрузки из Яндекс Календаря: %s", e)
            return []

    return _load_all_grant_and_kkt_events(days_ahead=days, min_year=min_year, force_local=True)


def _settings_schedule(block_id: str) -> None:
    try:
        from src.yandex_calendar_client import is_configured
        yandex_ok = is_configured()
    except Exception:
        yandex_ok = False
    source = _get("dashboard_schedule_source", "local")
    options = ["Локальный файл (grant_calendar.json)", "Яндекс Календарь (CalDAV)"]
    idx = 1 if source == "yandex" else 0
    choice = st.radio("Источник событий", options, index=idx, key=f"radio_source_{block_id}")
    st.session_state["dashboard_schedule_source"] = "yandex" if "Яндекс" in choice else "local"
    if "Яндекс" in choice and not yandex_ok:
        st.caption("Задайте YANDEX_CALENDAR_USER и YANDEX_CALENDAR_APP_PASSWORD в .env или в config (yandex_calendar). Пароль приложения: id.yandex.ru → Безопасность → Пароли приложений → Календарь.")
    days = st.number_input("Период (дней вперёд)", min_value=1, max_value=365, value=_get("dashboard_schedule_days", 30), key=f"ni_days_{block_id}")
    st.session_state["dashboard_schedule_days"] = days
    min_year = st.number_input("Показывать события не раньше года", min_value=2020, max_value=2030, value=_get("dashboard_schedule_min_year", 2026), key=f"ni_min_year_schedule_{block_id}")
    st.session_state["dashboard_schedule_min_year"] = min_year
    show_desc = st.checkbox("Показывать описание события", value=_get("dashboard_schedule_show_description", True), key=f"cb_desc_{block_id}")
    st.session_state["dashboard_schedule_show_description"] = show_desc
    path_override = st.text_input("Путь к локальному календарю (только для источника «Локальный файл»)", value=_get("dashboard_schedule_calendar_path", ""), key=f"path_cal_{block_id}")
    st.session_state["dashboard_schedule_calendar_path"] = path_override.strip()


def _content_schedule(block_id: str) -> None:
    days = _get("dashboard_schedule_days", 30)
    show_desc = _get("dashboard_schedule_show_description", True)
    source = _get("dashboard_schedule_source", "local")
    events = _load_schedule_events()
    if not events:
        if source == "yandex":
            msg = "Яндекс Календарь не настроен или за период событий нет. Задайте YANDEX_CALENDAR_USER и YANDEX_CALENDAR_APP_PASSWORD в .env или в настройках блока."
        else:
            msg = "Календарь гранта пуст. Добавьте data/grant_calendar.json или переключите источник на «Яндекс Календарь» в настройках."
        st.markdown(
            f'<div class="np-empty-state"><span class="np-empty-icon">📅</span><p>{msg}</p></div>',
            unsafe_allow_html=True,
        )
        return
    today = datetime.utcnow().date()
    end = today + timedelta(days=days)
    min_year = _get("dashboard_schedule_min_year", 2026)
    min_date = datetime(min_year, 1, 1).date()
    upcoming = []
    for ev in events:
        try:
            d = datetime.fromisoformat(ev["date"].replace("Z", "")).date()
        except (KeyError, ValueError):
            continue
        if d < min_date:
            continue
        if today <= d <= end:
            upcoming.append((d, ev.get("title", ""), ev.get("description", ""), ev.get("address", "")))
    upcoming.sort(key=lambda x: x[0])
    if not upcoming:
        st.markdown(
            f'<div class="np-empty-state"><span class="np-empty-icon">📅</span><p>Нет событий на ближайшие {days} дней</p></div>',
            unsafe_allow_html=True,
        )
        return
    if source == "yandex":
        st.caption("Источник: Яндекс Календарь (CalDAV)")
    for d, title, desc, address in upcoming:
        days_left = (d - today).days
        st.markdown(f"**{d}** ({days_left} дн.) — {title}")
        if show_desc and desc:
            st.caption(desc)
        if address and address.strip():
            from urllib.parse import quote
            maps_url = "https://yandex.ru/maps/?text=" + quote(address.strip())
            st.link_button("🗺️ Открыть в Яндекс.Картах", maps_url)


def block_schedule() -> None:
    _block_with_settings("📅 Ближайшие сроки по гранту", "schedule", _content_schedule, _settings_schedule)


# ---------- Блок: Напоминания ----------
def _settings_reminders(block_id: str) -> None:
    st.caption("Те же события, что в «Ближайшие сроки», с разметкой по срочности (за 7/3/1 дн., сегодня, просрочено).")
    min_year = st.number_input("Показывать события не раньше года", min_value=2020, max_value=2030, value=_get("dashboard_reminders_min_year", 2026), key=f"ni_reminders_min_year_{block_id}")
    st.session_state["dashboard_reminders_min_year"] = min_year
    st.caption("Email-напоминания отправляет scripts/grant_reminders.py по расписанию.")


def _content_reminders(block_id: str) -> None:
    events = _load_schedule_events()
    today = datetime.utcnow().date()
    min_year = _get("dashboard_reminders_min_year", 2026)
    cutoff = datetime(min_year, 1, 1).date()
    reminders_7 = []
    reminders_3 = []
    reminders_1 = []
    today_ev = []
    overdue = []
    for ev in events:
        try:
            d = datetime.fromisoformat(ev["date"].replace("Z", "")).date()
        except (KeyError, ValueError):
            continue
        if d < cutoff:
            continue
        days_left = (d - today).days
        title = ev.get("title", "Событие")
        if days_left < 0:
            overdue.append((d, title, days_left))
        elif days_left == 0:
            today_ev.append((d, title))
        elif days_left == 1:
            reminders_1.append((d, title))
        elif days_left == 3:
            reminders_3.append((d, title))
        elif days_left == 7:
            reminders_7.append((d, title))
    if not (reminders_7 or reminders_3 or reminders_1 or today_ev or overdue):
        st.markdown(
            '<div class="np-empty-state"><span class="np-empty-icon">🔔</span><p>Нет событий в зоне напоминаний (за 7/3/1 дн., сегодня, просрочено). Добавьте события в календарь.</p></div>',
            unsafe_allow_html=True,
        )
        return
    for d, title, days_over in overdue:
        st.markdown(f"🔴 **Просрочено** ({abs(days_over)} дн.) — {title} (*{d}*)")
    for d, title in today_ev:
        st.markdown(f"🔥 **Сегодня** — {title}")
    for d, title in reminders_1:
        st.markdown(f"❗ **За 1 дн.** — {title} (*{d}*)")
    for d, title in reminders_3:
        st.markdown(f"🔔 **За 3 дн.** — {title} (*{d}*)")
    for d, title in reminders_7:
        st.markdown(f"⚠️ **За 7 дн.** — {title} (*{d}*)")


def block_reminders() -> None:
    _block_with_settings("🔔 Напоминания", "reminders", _content_reminders, _settings_reminders)


# ---------- Блок: Ключевые контрольные точки (ККТ), АНО «Гранты Ямала» ----------
def _load_kkt() -> list[dict]:
    """Загружает ККТ из data/grant_kkt.json или из примера."""
    for path in (KKT_PATH, KKT_EXAMPLE):
        if path.exists():
            try:
                with open(path, encoding="utf-8") as f:
                    data = json.load(f)
                return data.get("points", data) if isinstance(data, dict) else data
            except Exception as e:
                logger.warning("Ошибка чтения ККТ %s: %s", path.name, e)
    return []


def _settings_kkt(block_id: str) -> None:
    min_year = st.number_input("Показывать ККТ с года", min_value=2020, max_value=2030, value=_get("dashboard_kkt_min_year", 2026), key=f"ni_kkt_min_year_{block_id}")
    st.session_state["dashboard_kkt_min_year"] = min_year
    st.caption(
        "ККТ создаются и редактируются в личном кабинете АНО «Гранты Ямала». "
        "На дашборде — локальная копия из data/grant_kkt.json."
    )


def _content_kkt(block_id: str) -> None:
    points = _load_kkt()
    if not points:
        st.info("Добавьте данные в data/grant_kkt.json (можно скопировать из data/grant_kkt.example.json).")
        return
    min_year = _get("dashboard_kkt_min_year", 2026)
    try:
        filtered = [p for p in points if not (p.get("date_end") or "").strip() or int(((p.get("date_end") or "").strip())[:4]) >= min_year]
    except ValueError:
        filtered = points
    if not filtered:
        st.caption(f"Нет ККТ с годом срока не раньше {min_year}. Измените фильтр в настройках блока.")
        return
    today = datetime.utcnow().date()

    def _esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    cards = []
    for p in filtered:
        date_end = (p.get("date_end") or "").strip()
        desc = _esc((p.get("description") or "").strip() or "—")
        status = p.get("status", "not_started")
        st_label = STATUS_LABELS.get(status, status)
        status_class = f"np-{status.replace(' ', '_')}"
        if date_end:
            try:
                d = datetime.strptime(date_end[:10], "%Y-%m-%d").date()
                if d < today and status != "done":
                    status_class = "np-overdue"
                    st_label = st_label + " (просрочено)"
            except ValueError:
                pass
        date_fmt = _esc(date_end) if date_end else "—"
        cards.append(
            f'<div class="np-milestone-card">'
            f'<span class="np-milestone-icon">🎯</span>'
            f'<div><div class="np-milestone-title">{desc}</div>'
            f'<div class="np-milestone-date">до {date_fmt}</div>'
            f'<div class="np-milestone-status {status_class}">{_esc(st_label)}</div></div></div>'
        )
    st.markdown('<div class="np-milestone-grid">' + "".join(cards) + "</div>", unsafe_allow_html=True)
    st.caption("Источник: АНО «Гранты Ямала». Редактирование — в личном кабинете информационной системы грантодателя.")


def block_kkt() -> None:
    _block_with_settings("🎯 Ключевые контрольные точки", "kkt", _content_kkt, _settings_kkt)


# Цвета фона для событий по этапам (Блок 1/2/3) в календаре
# Маппинг: ключ — подстрока для поиска в названии этапа (режим «содержит»)
STAGE_BG_COLORS = {
    "Подготовительный": "#e3f2fd",
    "Основной": "#e8f5e9",
    "Завершающий": "#fff3e0",
}


def _get_stage_bg_color(stage: str) -> str:
    """Возвращает цвет фона для этапа. Ищет подстроку в названии этапа (устойчиво к «Блок 1: Подготовительный» vs «Подготовительный»)."""
    if not stage:
        return "transparent"
    stage_lower = stage.lower()
    for key, color in STAGE_BG_COLORS.items():
        if key.lower() in stage_lower:
            return color
    return "transparent"


def _load_stage_ranges() -> list[tuple[str, str, str]]:
    """Возвращает список (stage_name, plan_start, plan_end) из grant_project_dashboard.json для привязки дат к этапам."""
    data = _load_grant_dashboard_data()
    stages = data.get("stages") or []
    out = []
    for s in stages:
        stage_name = (s.get("stage") or "").strip()
        start_s = (s.get("plan_start") or "").strip()[:10]
        end_s = (s.get("plan_end") or "").strip()[:10]
        if stage_name and (start_s or end_s):
            out.append((stage_name, start_s or end_s, end_s or start_s))
    return out


def _get_stage_for_date(d: datetime.date, stage_ranges: list[tuple[str, str, str]]) -> str:
    """Определяет этап по дате: дата попадает в plan_start..plan_end. Возвращает название этапа или пустую строку."""
    d_str = d.isoformat()[:10]
    for stage_name, start_s, end_s in stage_ranges:
        if not start_s or not end_s:
            continue
        if start_s <= d_str <= end_s:
            return stage_name
    return ""


def _load_all_grant_and_kkt_events(
    days_ahead: int = 365,
    min_year: int = 2026,
    start_date: datetime.date | None = None,
    end_date: datetime.date | None = None,
    force_local: bool = False,
) -> list[dict]:
    """Объединяет события из календаря гранта, ККТ и этапов (grant_project_dashboard). У каждого события поле stage для раскраски по этапу. Если заданы start_date/end_date — ими ограничиваем период вместо today..today+days_ahead. force_local=True — всегда читать из grant_calendar.json, игнорируя настройку источника «Яндекс Календарь»."""
    today = datetime.utcnow().date()
    if start_date is not None and end_date is not None:
        period_start, period_end = start_date, end_date
    else:
        period_start = today
        period_end = today + timedelta(days=days_ahead)
    cutoff = datetime(min_year, 1, 1).date()
    stage_ranges = _load_stage_ranges()
    out = []

    grant_events = _load_calendar_local() if force_local else _load_schedule_events()
    for ev in grant_events:
        try:
            d = datetime.fromisoformat(ev["date"].replace("Z", "")).date()
        except (KeyError, ValueError):
            continue
        if d < cutoff:
            continue
        if period_start <= d <= period_end:
            out.append({
                "date": d.isoformat(),
                "title": ev.get("title", "Событие"),
                "description": ev.get("description", ""),
                "address": ev.get("address", ""),
                "source": "grant",
                "stage": _get_stage_for_date(d, stage_ranges),
            })
    for p in _load_kkt():
        date_end = (p.get("date_end") or "").strip()
        if not date_end:
            continue
        try:
            d = datetime.strptime(date_end[:10], "%Y-%m-%d").date()
        except ValueError:
            continue
        if d < cutoff:
            continue
        if period_start <= d <= period_end:
            desc = p.get("description", "")
            expected = (p.get("expected_result") or "").strip()
            out.append({
                "date": d.isoformat(),
                "title": f"ККТ: {desc}" if desc else "ККТ",
                "description": f"Ожидаемый результат: {expected}" if expected else "",
                "address": "",
                "source": "kkt",
                "stage": _get_stage_for_date(d, stage_ranges),
            })
    # События из этапов (активности с plan_start/plan_end — одна дата на активность: plan_end или plan_start)
    data = _load_grant_dashboard_data()
    for s in (data.get("stages") or []):
        stage_name = (s.get("stage") or "").strip()
        activity = (s.get("activity") or "").strip()
        date_s = (s.get("plan_end") or s.get("plan_start") or "").strip()[:10]
        if not date_s or not activity:
            continue
        try:
            d = datetime.strptime(date_s, "%Y-%m-%d").date()
        except ValueError:
            continue
        if d < cutoff or d < period_start or d > period_end:
            continue
        out.append({
            "date": d.isoformat(),
            "title": activity,
            "description": f"Этап: {stage_name}",
            "address": "",
            "source": "stage",
            "stage": stage_name,
        })
    # События, импортированные из календаря Нейропульс (двусторонняя синхронизация)
    if MANUAL_CALENDAR_EVENTS_PATH.exists():
        try:
            with open(MANUAL_CALENDAR_EVENTS_PATH, encoding="utf-8") as f:
                manual_list = json.load(f)
            for ev in (manual_list if isinstance(manual_list, list) else []):
                date_s = (ev.get("date") or "").strip()[:10]
                if not date_s:
                    continue
                try:
                    d = datetime.strptime(date_s, "%Y-%m-%d").date()
                except ValueError:
                    continue
                if d < cutoff or d < period_start or d > period_end:
                    continue
                out.append({
                    "date": d.isoformat(),
                    "title": (ev.get("title") or "Событие").strip(),
                    "description": (ev.get("description") or "").strip(),
                    "address": (ev.get("address") or "").strip(),
                    "source": "manual",
                    "stage": _get_stage_for_date(d, stage_ranges),
                })
        except Exception as e:
            logger.warning("Ошибка чтения %s: %s", MANUAL_CALENDAR_EVENTS_PATH.name, e)
    out.sort(key=lambda x: (x["date"], x["title"]))
    return out


# ---------- Блок 2: Аудит чата ----------
def _settings_audit(block_id: str) -> None:
    days = st.number_input("Период (дней назад)", min_value=1, max_value=365, value=_get("dashboard_audit_days", 7), key=f"ni_audit_days_{block_id}")
    st.session_state["dashboard_audit_days"] = days
    group = st.radio("Группировка", ["По дням", "По сессиям"], index=0 if _get("dashboard_audit_group", "day") == "day" else 1, key=f"radio_audit_{block_id}")
    st.session_state["dashboard_audit_group"] = "day" if "дням" in group else "session"
    limit = st.number_input("Макс. записей в таблице", min_value=10, max_value=1000, value=_get("dashboard_audit_limit", 100), key=f"ni_audit_lim_{block_id}")
    st.session_state["dashboard_audit_limit"] = limit


def _content_audit(block_id: str) -> None:
    days = _get("dashboard_audit_days", 7)
    limit = _get("dashboard_audit_limit", 100)
    group = _get("dashboard_audit_group", "day")
    if not AUDIT_LOG.exists():
        st.caption("Файл аудита чата пока пуст (data/chat_audit.jsonl).")
        return
    records = []
    try:
        with open(AUDIT_LOG, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    records.append(json.loads(line))
                except json.JSONDecodeError:
                    continue
    except Exception as e:
        st.error(f"Ошибка чтения аудита: {e}")
        return
    if not records:
        st.caption("Записей в аудите нет.")
        return
    cutoff = (datetime.utcnow() - timedelta(days=days)).isoformat() + "Z"
    records = [r for r in records if (r.get("ts") or "") >= cutoff][-limit:]
    if not records:
        st.caption(f"За последние {days} дн. записей нет.")
        return
    if group == "day":
        by_day = {}
        for r in records:
            ts = (r.get("ts") or "")[:10]
            by_day.setdefault(ts, {"count": 0, "prompt_len": 0, "reply_len": 0})
            by_day[ts]["count"] += 1
            by_day[ts]["prompt_len"] += r.get("prompt_len", 0)
            by_day[ts]["reply_len"] += r.get("reply_len", 0)
        rows = [{"Дата": k, "Запросов": v["count"], "Символов запрос": v["prompt_len"], "Символов ответ": v["reply_len"]} for k, v in sorted(by_day.items(), reverse=True)]
    else:
        rows = [{"Время": r.get("ts", ""), "Хеш сессии": r.get("conv_id_hash", ""), "Длина запроса": r.get("prompt_len", 0), "Длина ответа": r.get("reply_len", 0)} for r in reversed(records)]
    st.dataframe(rows, width="stretch", hide_index=True)


def block_audit() -> None:
    _block_with_settings("📊 Аудит чата", "audit", _content_audit, _settings_audit)


# ---------- Блок 3: Статус агента и API ----------
def _settings_status(block_id: str) -> None:
    check_vs = st.checkbox("Проверять доступность Vector Store", value=_get("dashboard_status_check_vs", True), key=f"cb_vs_{block_id}")
    st.session_state["dashboard_status_check_vs"] = check_vs


def _content_status(block_id: str) -> None:
    config = load_config()
    project_cfg = config.get("project", {})
    ya_cfg = config.get("yandex_ai_studio", {})
    agent_name = ya_cfg.get("agent_name") or "—"
    version = project_cfg.get("version") or "—"

    try:
        from src.yandex_ai_client import get_yandex_ai_config
        ya = get_yandex_ai_config()
        has_key = bool((ya.get("api_key") or "").strip())
        folder_id = (ya.get("folder_id") or "").strip()
        model_spec = ya.get("agent_id") or ya.get("model") or "—"
    except Exception:
        has_key = False
        folder_id = ""
        model_spec = "—"

    st.markdown(f"**Проект:** {project_cfg.get('name', '—')} v{version}")
    st.markdown(f"**Агент:** {agent_name}")
    st.markdown(f"**Модель/агент в запросах:** {model_spec}")
    st.caption(f"API ключ: {'задан' if has_key else 'не задан'} | Folder ID: {'задан' if folder_id else 'не задан'}")

    configured = has_key and folder_id
    if not configured:
        st.warning("Настройте YANDEX_API_KEY и YANDEX_FOLDER_ID в .env (см. .env.example) для проверки API.")
        return

    try:
        from src.yandex_ai_client import get_client
        get_client()
        st.success("Подключение к Yandex AI Studio: OK")
    except Exception as e:
        st.error(f"Ошибка подключения: {e}")
        return

    if _get("dashboard_status_check_vs", True):
        try:
            from src.vector_store_client import list_vector_stores
            r = list_vector_stores(limit=5)
            items = getattr(r, "data", r) if not isinstance(r, list) else r or []
            st.caption(f"Vector Store: найдено индексов: {len(items)}")
        except Exception as e:
            err_str = str(e).lower()
            if "403" in err_str and ("permission" in err_str or "folder" in err_str):
                st.caption("Vector Store: 403. Проверьте, что в config/.env указан **YANDEX_FOLDER_ID** — идентификатор **каталога** (обычно b1g...), а не идентификатор агента.")
            else:
                st.caption(f"Vector Store: ошибка — {e}")


def block_status() -> None:
    _block_with_settings("🟢 Статус агента и API", "status", _content_status, _settings_status)


# ---------- Блок 4: Быстрые действия (grant_tools) — кнопки отправляют формулировки в диалог с агентом ----------
def _settings_tools(block_id: str) -> None:
    sort_by = st.radio("Сортировка", ["По id", "По названию"], index=1 if _get("dashboard_tools_sort", "title") == "title" else 0, key=f"radio_tools_{block_id}")
    st.session_state["dashboard_tools_sort"] = "title" if "названию" in sort_by else "id"


def _content_tools(block_id: str) -> None:
    config = load_config()
    tools = config.get("grant_tools", [])
    sort = _get("dashboard_tools_sort", "title")
    if sort == "title":
        tools = sorted(tools, key=lambda t: (t.get("title") or t.get("id") or ""))
    else:
        tools = sorted(tools, key=lambda t: (t.get("id") or t.get("title") or ""))
    if not tools:
        st.caption("В config нет grant_tools. Добавьте кнопки в config.json.")
        return
    st.caption("Нажмите — формулировка уйдёт в диалог слева.")
    for t in tools:
        title = t.get("title") or t.get("id") or "—"
        prompt = (t.get("prompt") or "").strip()
        if st.button(title, key=f"tool_btn_{t.get('id', title)}", width="stretch"):
            if "dashboard_prompt_to_send" not in st.session_state:
                st.session_state.dashboard_prompt_to_send = []
            st.session_state.dashboard_prompt_to_send.append(prompt)
            st.rerun()


def block_tools() -> None:
    _block_with_settings("🔘 Быстрые действия по гранту", "tools", _content_tools, _settings_tools)


# ---------- Блок 5: База знаний (Vector Store) ----------
def _settings_vs(block_id: str) -> None:
    limit = st.number_input("Макс. файлов для сводки", min_value=5, max_value=200, value=_get("dashboard_vs_files_limit", 50), key=f"ni_vs_lim_{block_id}")
    st.session_state["dashboard_vs_files_limit"] = limit
    st.caption("Сводка строится по первым N файлам индекса.")


def _content_vs(block_id: str) -> None:
    config = load_config()
    vs_id = (
        os.environ.get("YANDEX_VECTOR_STORE_ID", "").strip()
        or (config.get("yandex_ai_studio", {}) or {}).get("vector_store_id") or ""
    ).strip()
    if not vs_id:
        st.info(
            "Задайте **vector_store_id**: в `config/config.json` — `yandex_ai_studio.vector_store_id` "
            "или в `config/.env` — `YANDEX_VECTOR_STORE_ID` (ID поискового индекса в Yandex AI Studio)."
        )
        return
    limit = _get("dashboard_vs_files_limit", 20)
    try:
        from src.vector_store_client import get_vector_store, list_vector_store_files
        vs = get_vector_store(vs_id)
        name = getattr(vs, "name", None) or vs_id
        r = list_vector_store_files(vs_id, limit=limit, status=None)
        items = getattr(r, "data", r) if not isinstance(r, list) else r or []
        # Сводные данные
        total_bytes = 0
        by_status = {}
        for f in items:
            size = getattr(f, "bytes", None)
            if size is not None:
                total_bytes += int(size)
            stt = getattr(f, "status", None) or "—"
            by_status[stt] = by_status.get(stt, 0) + 1
        def _fmt_size(b: int) -> str:
            if b < 1024:
                return f"{b} B"
            if b < 1024 * 1024:
                return f"{b / 1024:.1f} KB"
            return f"{b / (1024 * 1024):.1f} MB"
        st.caption(f"**Индекс:** {name}")
        st.metric("Файлов", len(items))
        st.metric("Общий размер", _fmt_size(total_bytes) if total_bytes else "—")
        if by_status:
            st.caption("По статусам: " + ", ".join(f"{s}: {n}" for s, n in sorted(by_status.items())))
        if not items:
            st.caption("В индексе пока нет файлов.")
    except Exception as e:
        st.error(f"Ошибка Vector Store: {e}")


def block_vector_store() -> None:
    _block_with_settings("📚 База знаний (Vector Store)", "vector_store", _content_vs, _settings_vs)


# ---------- Блок: Агент Парсер ----------
PARSER_LAST_RUN_PATH = PROJECT_ROOT / "data" / "parser_last_run.json"


def _load_parser_last_run() -> dict | None:
    if not PARSER_LAST_RUN_PATH.exists():
        return None
    try:
        with open(PARSER_LAST_RUN_PATH, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _save_parser_last_run(result: dict) -> None:
    PARSER_LAST_RUN_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(PARSER_LAST_RUN_PATH, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)


def _save_parser_sources(channels: list[str], disk_links: list[str]) -> tuple[bool, str]:
    """Сохраняет каналы и ссылки в config/config.json → parser."""
    try:
        from src.agent_api_client import CONFIG_PATH, load_config
        config = {}
        cfg_path = CONFIG_PATH
        if cfg_path.exists():
            with open(cfg_path, encoding="utf-8") as f:
                config = json.load(f)
        else:
            # Копируем из example если нет config.json
            from src.agent_api_client import CONFIG_EXAMPLE_PATH
            if CONFIG_EXAMPLE_PATH.exists():
                with open(CONFIG_EXAMPLE_PATH, encoding="utf-8") as f:
                    config = json.load(f)
        if "parser" not in config:
            config["parser"] = {}
        config["parser"]["telegram_channels"] = channels
        config["parser"]["disk_public_links"] = disk_links
        cfg_path.parent.mkdir(parents=True, exist_ok=True)
        with open(cfg_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        return True, f"Сохранено: {len(channels)} канал(ов) Telegram, {len(disk_links)} ссылок Диска."
    except Exception as e:
        return False, str(e)


def _load_parser_raw_sources() -> tuple[list[str], list[str]]:
    """Читает сырые строки telegram_channels и disk_public_links из config.json (с сохранёнными названиями)."""
    try:
        from src.agent_api_client import CONFIG_PATH, load_config
        path = CONFIG_PATH
        if path.exists():
            with open(path, encoding="utf-8") as f:
                config = json.load(f)
        else:
            config = load_config()
        p = config.get("parser") or {}
        tg = p.get("telegram_channels") or []
        disk = p.get("disk_public_links") or []
        if isinstance(tg, str):
            tg = [s.strip() for s in tg.split(",") if s.strip()]
        if isinstance(disk, str):
            disk = [s.strip() for s in disk.split(",") if s.strip()]
        return list(tg), list(disk)
    except Exception:
        return [], []


def _settings_parser(block_id: str) -> None:
    channels_raw, disk_raw = _load_parser_raw_sources()

    # ── Telegram-каналы ────────────────────────────────────────
    st.markdown("#### 📱 Telegram-каналы")
    st.caption(
        "Один канал на строку. После ссылки можно добавить название через пробел:\n\n"
        "`@channel` · `https://t.me/channel Название` · `https://t.me/c/ID Описание`"
    )

    # Превью текущих каналов в виде таблицы
    if channels_raw:
        rows_html = ""
        for line in channels_raw:
            parts = line.split(None, 1)
            url = parts[0] if parts else line
            name = parts[1] if len(parts) > 1 else "—"
            rows_html += (
                f'<tr><td style="font-family:monospace;font-size:12px;padding:4px 8px;'
                f'color:#1565c0;word-break:break-all">{url}</td>'
                f'<td style="padding:4px 8px;font-size:12px;color:#475569">{name}</td></tr>'
            )
        st.markdown(
            f'<table style="width:100%;border-collapse:collapse;border:1px solid #e2e8f0;border-radius:8px;overflow:hidden;margin-bottom:6px">'
            f'<thead><tr><th style="text-align:left;padding:5px 8px;background:#f8fafc;font-size:11px;color:#64748b">Ссылка / канал</th>'
            f'<th style="text-align:left;padding:5px 8px;background:#f8fafc;font-size:11px;color:#64748b">Название</th></tr></thead>'
            f'<tbody>{rows_html}</tbody></table>',
            unsafe_allow_html=True,
        )

    tg_text = st.text_area(
        "Каналы Telegram",
        value="\n".join(channels_raw),
        height=130,
        placeholder="@mychannel\nhttps://t.me/other_channel НКО Победители\nhttps://t.me/c/3847911347 Гранты ЯНАО",
        key=f"parser_tg_input_{block_id}",
        label_visibility="collapsed",
    )

    # ── Ссылки Яндекс.Диска ────────────────────────────────────
    st.markdown("#### 📁 Ссылки Яндекс.Диска")
    st.caption("Один URL на строку. Можно добавить название: `https://disk.yandex.ru/d/xxx Документы проекта`")

    if disk_raw:
        rows_html = ""
        for line in disk_raw:
            parts = line.split(None, 1)
            url = parts[0] if parts else line
            name = parts[1] if len(parts) > 1 else "—"
            rows_html += (
                f'<tr><td style="font-family:monospace;font-size:12px;padding:4px 8px;'
                f'color:#1565c0;word-break:break-all">{url}</td>'
                f'<td style="padding:4px 8px;font-size:12px;color:#475569">{name}</td></tr>'
            )
        st.markdown(
            f'<table style="width:100%;border-collapse:collapse;border:1px solid #e2e8f0;border-radius:8px;overflow:hidden;margin-bottom:6px">'
            f'<thead><tr><th style="text-align:left;padding:5px 8px;background:#f8fafc;font-size:11px;color:#64748b">Ссылка</th>'
            f'<th style="text-align:left;padding:5px 8px;background:#f8fafc;font-size:11px;color:#64748b">Название</th></tr></thead>'
            f'<tbody>{rows_html}</tbody></table>',
            unsafe_allow_html=True,
        )

    disk_text = st.text_area(
        "Ссылки Диска",
        value="\n".join(disk_raw),
        height=90,
        placeholder="https://disk.yandex.ru/d/xxxxx Документы гранта",
        key=f"parser_disk_input_{block_id}",
        label_visibility="collapsed",
    )

    if st.button("💾 Сохранить источники", key=f"parser_save_sources_{block_id}", type="primary"):
        new_channels = [s.strip() for s in tg_text.splitlines() if s.strip()]
        new_disk = [s.strip() for s in disk_text.splitlines() if s.strip()]
        ok, msg = _save_parser_sources(new_channels, new_disk)
        if ok:
            st.success(msg)
            st.rerun()
        else:
            st.error(f"Ошибка сохранения: {msg}")

    st.divider()
    st.markdown("**Секреты Telegram** (для приватных каналов) — задайте в `config/.env`:")
    st.code("TELEGRAM_API_ID=...\nTELEGRAM_API_HASH=...\nTELEGRAM_PHONE=+7...", language="bash")
    st.caption("Получить API ID и Hash: https://my.telegram.org/apps")


def _content_parser(block_id: str) -> None:
    try:
        from src.parser import get_parser_config, run_parser_pipeline
    except ImportError as e:
        st.warning(f"Модуль парсера недоступен: {e}")
        return
    cfg = get_parser_config()
    vs_id = (cfg.get("vector_store_id") or "").strip()
    channels = cfg.get("telegram_channels") or []
    disk_links = cfg.get("disk_public_links") or []
    has_tg_creds = bool((cfg.get("telegram_api_id") or "").strip() and (cfg.get("telegram_api_hash") or "").strip())

    st.caption(f"Каналов Telegram: {len(channels)}, ссылок Яндекс.Диск: {len(disk_links)}. Vector Store: {'задан' if vs_id else 'не задан'}.")
    hints = []
    if channels and not has_tg_creds:
        hints.append("Для парсинга Telegram задайте в .env: TELEGRAM_API_ID, TELEGRAM_API_HASH, TELEGRAM_PHONE.")
    if not channels and not disk_links:
        hints.append("Добавьте каналы в parser.telegram_channels и/или ссылки в parser.disk_public_links в config/config.json.")
    if (channels or disk_links) and not vs_id:
        hints.append("Укажите vector_store_id в parser или YANDEX_VECTOR_STORE_ID в .env — иначе загрузка в индекс не выполнится.")
    for h in hints:
        st.warning(h)

    run_key = "parser_run_clicked"
    if run_key not in st.session_state:
        st.session_state[run_key] = None

    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("📱 Парсинг Telegram", key="btn_parser_tg"):
            st.session_state[run_key] = ("telegram", True, False)
            st.rerun()
    with col2:
        if st.button("📁 Парсинг Яндекс.Диск", key="btn_parser_disk"):
            st.session_state[run_key] = ("disk", False, True)
            st.rerun()
    with col3:
        if st.button("🔄 Полный парсинг", key="btn_parser_full"):
            st.session_state[run_key] = ("full", True, True)
            st.rerun()

    if st.session_state[run_key]:
        _mode, run_tg, run_disk = st.session_state[run_key]
        st.session_state[run_key] = None
        with st.spinner("Запуск парсера…"):
            result = run_parser_pipeline(run_telegram=run_tg, run_disk=run_disk)
        _save_parser_last_run(result)
        if result.get("ok"):
            st.success(result.get("details", "Готово."))
        else:
            st.error("; ".join(result.get("errors", ["Ошибка"])))
        st.caption(f"Загружено в индекс: {result.get('files_uploaded', 0)} файлов.")

    last = _load_parser_last_run()
    if last:
        st.caption(f"Последний запуск: {last.get('details', '—')}")


def block_parser() -> None:
    _block_with_settings("🔀 Агент Парсер", "parser", _content_parser, _settings_parser)


# ---------- Блок: Управление рисками ----------
def _load_risks() -> list[dict]:
    for path in (RISKS_PATH, RISKS_EXAMPLE):
        if path.exists():
            try:
                with open(path, encoding="utf-8") as f:
                    data = json.load(f)
                return data if isinstance(data, list) else data.get("risks", [])
            except Exception as e:
                logger.warning("Ошибка чтения рисков %s: %s", path.name, e)
    return []


def _settings_risks(block_id: str) -> None:
    st.caption("Данные из data/grant_risks.json (скопируйте из grant_risks.example.json). Редактирование — вручную в файле или добавьте форму в настройках.")


def _content_risks(block_id: str) -> None:
    items = _load_risks()
    if not items:
        st.info("Добавьте данные в data/grant_risks.json (пример: data/grant_risks.example.json).")
        return
    rows = []
    for r in items:
        rows.append({
            "Описание": r.get("description", ""),
            "Вероятность": r.get("probability", ""),
            "Влияние": r.get("impact", ""),
            "Митигация": r.get("mitigation", ""),
            "Статус": r.get("status", ""),
        })
    st.dataframe(rows, width="stretch", hide_index=True, column_config={"Описание": st.column_config.TextColumn("Описание", width="medium"), "Митигация": st.column_config.TextColumn("Митигация", width="medium")})


def block_risks() -> None:
    _block_with_settings("⚠️ Управление рисками", "risks", _content_risks, _settings_risks)


# ---------- Блок: Статистика агента ----------
def _settings_agent_stats(block_id: str) -> None:
    days = st.number_input("Период (дней назад)", min_value=1, max_value=365, value=_get("dashboard_agent_stats_days", 30), key=f"ni_agent_stats_days_{block_id}")
    st.session_state["dashboard_agent_stats_days"] = days


def _content_agent_stats(block_id: str) -> None:
    days = _get("dashboard_agent_stats_days", 30)
    cutoff = (datetime.utcnow() - timedelta(days=days)).isoformat() + "Z"
    if not AUDIT_LOG.exists():
        st.caption("Файл аудита чата пуст (data/chat_audit.jsonl).")
        return
    records = []
    try:
        with open(AUDIT_LOG, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    r = json.loads(line)
                    if (r.get("ts") or "") >= cutoff:
                        records.append(r)
                except json.JSONDecodeError:
                    continue
    except Exception as e:
        st.error(f"Ошибка чтения аудита: {e}")
        return
    if not records:
        st.caption(f"За последние {days} дн. записей нет.")
        return
    by_day: dict[str, dict[str, Any]] = {}
    for r in records:
        ts = (r.get("ts") or "")[:10]
        if ts not in by_day:
            by_day[ts] = {"count": 0, "prompt_len": 0, "reply_len": 0}
        by_day[ts]["count"] += 1
        by_day[ts]["prompt_len"] += r.get("prompt_len", 0)
        by_day[ts]["reply_len"] += r.get("reply_len", 0)
    total_requests = sum(d["count"] for d in by_day.values())
    total_prompt = sum(d["prompt_len"] for d in by_day.values())
    total_reply = sum(d["reply_len"] for d in by_day.values())
    st.metric("Запросов к агенту", total_requests)
    st.metric("Символов (запросы)", total_prompt)
    st.metric("Символов (ответы)", total_reply)
    sorted_days = sorted(by_day.items(), key=lambda x: x[0])
    if sorted_days:
        chart_data = [{"Дата": d, "Запросов": v["count"]} for d, v in sorted_days]
        st.bar_chart(chart_data, x="Дата", y="Запросов", height=200)


def block_agent_stats() -> None:
    _block_with_settings("📈 Статистика агента", "agent_stats", _content_agent_stats, _settings_agent_stats)


# ---------- Блок: Настройки уведомлений ----------
def _load_notification_settings() -> dict:
    for path in (NOTIFICATION_SETTINGS_PATH, NOTIFICATION_SETTINGS_EXAMPLE):
        if path.exists():
            try:
                with open(path, encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                logger.warning("Ошибка чтения notification_settings %s: %s", path.name, e)
    return {"email": "", "reminders_enabled": True, "days_before": [7, 3, 1]}


def _save_notification_settings(data: dict) -> bool:
    try:
        NOTIFICATION_SETTINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
        with open(NOTIFICATION_SETTINGS_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception as e:
        logger.warning("Ошибка сохранения notification_settings: %s", e)
        return False


def _settings_notifications(block_id: str) -> None:
    data = _load_notification_settings()
    email = st.text_input("Email для напоминаний", value=data.get("email", ""), key=f"notif_email_{block_id}")
    enabled = st.checkbox("Включить напоминания по срокам", value=data.get("reminders_enabled", True), key=f"notif_enabled_{block_id}")
    if st.button("Сохранить настройки", key=f"notif_save_{block_id}"):
        new_data = {"email": email.strip(), "reminders_enabled": enabled, "days_before": data.get("days_before", [7, 3, 1])}
        if _save_notification_settings(new_data):
            st.success("Сохранено. scripts/grant_reminders.py при запуске может учитывать email из data/notification_settings.json.")
        else:
            st.error("Ошибка сохранения.")
    st.caption("Получатель по умолчанию также задаётся в config/.env: NOTIFICATION_EMAIL. Отправка писем — scripts/grant_reminders.py.")


def _content_notifications(block_id: str) -> None:
    data = _load_notification_settings()
    st.caption(f"Email: {data.get('email') or '(не задан, используется NOTIFICATION_EMAIL из .env)'}")
    st.caption(f"Напоминания: {'включены' if data.get('reminders_enabled', True) else 'выключены'}.")


def block_notifications() -> None:
    _block_with_settings("🔔 Настройки уведомлений", "notifications", _content_notifications, _settings_notifications)


# ---------- Блок: Коммуникации с грантодателем ----------
def _load_communications() -> list[dict]:
    for path in (COMMUNICATIONS_PATH, COMMUNICATIONS_EXAMPLE):
        if path.exists():
            try:
                with open(path, encoding="utf-8") as f:
                    data = json.load(f)
                return data if isinstance(data, list) else data.get("items", [])
            except Exception as e:
                logger.warning("Ошибка чтения communications %s: %s", path.name, e)
    return []


def _settings_communications(block_id: str) -> None:
    st.caption("Данные из data/grant_communications.json (пример: grant_communications.example.json).")


def _content_communications(block_id: str) -> None:
    items = _load_communications()
    if not items:
        st.info("Добавьте данные в data/grant_communications.json.")
        return
    rows = [{"Дата": c.get("date", ""), "Тема": c.get("topic", ""), "Тип": c.get("type", ""), "Статус": c.get("status", "")} for c in items]
    st.dataframe(rows, width="stretch", hide_index=True)


def block_communications() -> None:
    _block_with_settings("📧 Коммуникации с грантодателем", "communications", _content_communications, _settings_communications)


# ---------- Блок: Файловый архив ----------
def _load_documents_archive() -> list[dict]:
    for path in (DOCUMENTS_ARCHIVE_PATH, DOCUMENTS_ARCHIVE_EXAMPLE):
        if path.exists():
            try:
                with open(path, encoding="utf-8") as f:
                    data = json.load(f)
                return data if isinstance(data, list) else data.get("documents", [])
            except Exception as e:
                logger.warning("Ошибка чтения documents_archive %s: %s", path.name, e)
    return []


def _settings_documents_archive(block_id: str) -> None:
    st.caption("Данные из data/grant_documents_archive.json (пример: grant_documents_archive.example.json).")


def _content_documents_archive(block_id: str) -> None:
    items = _load_documents_archive()
    if not items:
        st.info("Добавьте данные в data/grant_documents_archive.json.")
        return
    rows = [{"Название": d.get("name", ""), "Тип": d.get("type", ""), "Статус": d.get("status", ""), "Дата": d.get("date", "")} for d in items]
    st.dataframe(rows, width="stretch", hide_index=True)
    for i, d in enumerate(items):
        link = (d.get("link") or "").strip()
        if link:
            st.link_button(f"Открыть: {(d.get('name') or '')[:40]}", link, key=f"archive_link_{block_id}_{i}")


def block_documents_archive() -> None:
    _block_with_settings("📁 Файловый архив", "documents_archive", _content_documents_archive, _settings_documents_archive)


# ---------- Блок: Календарь (виджет месяц) ----------
def _settings_calendar_month(block_id: str) -> None:
    st.caption("Сетка месяца. Дни с событиями выделены. Источник событий — тот же, что в «Ближайшие сроки».")
    year = st.number_input("Год", min_value=2024, max_value=2030, value=_get("dashboard_calendar_month_year", datetime.utcnow().year), key=f"ni_cal_year_{block_id}")
    st.session_state["dashboard_calendar_month_year"] = year
    month = st.number_input("Месяц", min_value=1, max_value=12, value=_get("dashboard_calendar_month_month", datetime.utcnow().month), key=f"ni_cal_month_{block_id}")
    st.session_state["dashboard_calendar_month_month"] = month
    min_year = st.number_input("События не раньше года", min_value=2020, max_value=2030, value=_get("dashboard_calendar_month_min_year", 2026), key=f"ni_cal_min_year_{block_id}")
    st.session_state["dashboard_calendar_month_min_year"] = min_year


def _content_calendar_month(block_id: str) -> None:
    """Календарь месяца с раскраской дней по этапам (Блок 1 — голубой, Блок 2 — зелёный, Блок 3 — оранжевый)."""
    import calendar as cal_mod
    year = _get("dashboard_calendar_month_year", datetime.utcnow().year)
    month = _get("dashboard_calendar_month_month", datetime.utcnow().month)
    min_year = _get("dashboard_calendar_month_min_year", 2026)
    first = datetime(year, month, 1).date()
    last_day = cal_mod.monthrange(year, month)[1]
    last = datetime(year, month, last_day).date()
    events = _load_all_grant_and_kkt_events(
        min_year=min_year,
        start_date=first,
        end_date=last,
        force_local=True,
    )
    min_date = datetime(min_year, 1, 1).date()
    import calendar
    cal = calendar.Calendar(firstweekday=0)
    weeks = cal.monthdayscalendar(year, month)
    events_by_day: dict[int, list[dict]] = {}
    for ev in events:
        try:
            d = datetime.fromisoformat(ev["date"].replace("Z", "")).date()
        except (KeyError, ValueError):
            continue
        if d < min_date:
            continue
        if d.year == year and d.month == month:
            events_by_day.setdefault(d.day, []).append(ev)
    weekdays = "Пн Вт Ср Чт Пт Сб Вс".split()
    st.caption(f"**{year}, {month}**")
    # Таблица: цвет фона ячейки по этапу первого события в этот день
    cells_html = []
    for week in weeks:
        for day in week:
            if day == 0:
                cells_html.append('<td class="cal-empty"></td>')
            else:
                evs = events_by_day.get(day, [])
                stage = evs[0].get("stage", "") if evs else ""
                bg = _get_stage_bg_color(stage)
                mark = "•" if evs else ""
                style = f"background-color:{bg};" if bg != "transparent" else ""
                cells_html.append(f'<td class="cal-day" style="{style}">{day}{mark}</td>')
    rows = []
    for i in range(0, len(cells_html), 7):
        rows.append("<tr>" + "".join(cells_html[i : i + 7]) + "</tr>")
    table_html = "<table class=\"cal-table\"><thead><tr>" + "".join(f"<th>{w}</th>" for w in weekdays) + "</tr></thead><tbody>" + "".join(rows) + "</tbody></table>"
    st.markdown(
        f'<style>.cal-table {{ border-collapse: collapse; font-size: 0.9rem; }} .cal-table td, .cal-table th {{ border: 1px solid #ddd; padding: 0.35rem 0.5rem; text-align: center; }} .cal-empty {{ background: #fafafa; }} .cal-day {{ border-radius: 4px; }}</style>{table_html}',
        unsafe_allow_html=True,
    )
    st.caption("• — день с событием. Цвет: Блок 1 — голубой, Блок 2 — зелёный, Блок 3 — оранжевый. Список — в «Ближайшие сроки» и «Календарь Нейропульс».")


def block_calendar_month() -> None:
    _block_with_settings("📅 Календарь (месяц)", "calendar_month", _content_calendar_month, _settings_calendar_month)


# ---------- Блок: Календарь Нейропульс (отдельный фрейм) ----------
def _settings_neuropulse_cal(block_id: str) -> None:
    days = st.number_input("Период (дней вперёд)", min_value=1, max_value=365, value=_get("dashboard_neuropulse_days", 60), key=f"ni_neuropulse_days_{block_id}")
    st.session_state["dashboard_neuropulse_days"] = days
    min_year = st.number_input("События не раньше года", min_value=2020, max_value=2030, value=_get("dashboard_neuropulse_min_year", 2026), key=f"ni_neuropulse_min_year_{block_id}")
    st.session_state["dashboard_neuropulse_min_year"] = min_year
    show_desc = st.checkbox("Показывать описание", value=_get("dashboard_neuropulse_show_desc", True), key=f"cb_neuropulse_desc_{block_id}")
    st.session_state["dashboard_neuropulse_show_desc"] = show_desc
    auto_sync = st.checkbox("Синхронизировать автоматически при открытии блока", value=_get("dashboard_neuropulse_auto_sync", True), key=f"cb_neuropulse_auto_sync_{block_id}", help="События гранта и ККТ выгружаются в Календарь Нейропульс при открытии дашборда (без дубликатов).")
    st.session_state["dashboard_neuropulse_auto_sync"] = auto_sync
    st.caption("Показываются все события гранта и ККТ с метками [Грант] / [ККТ]. Виджет — календарь Яндекса. После синхронизации события отображаются в сетке календаря.")
    st.divider()
    st.caption("**Синхронизация:** события гранта (grant_calendar.json), ККТ (grant_kkt.json) и даты этапов (grant_project_dashboard.json).")
    if st.button("Очистить кэш отметок ✓", key=f"btn_clear_marks_cache_{block_id}", help="Удаляет локально сохранённые отметки и пересобирает их только по фактическому состоянию календаря."):
        try:
            from src.yandex_calendar_client import clear_persisted_event_keys
            clear_persisted_event_keys()
            for key in list(st.session_state.keys()):
                if str(key).startswith("np_add_cal_ok_"):
                    del st.session_state[key]
            st.success("Кэш отметок очищен. После перезагрузки блока ✓ будут построены заново по данным календаря.")
            st.rerun()
        except Exception as e:
            st.error(f"Ошибка очистки кэша: {e}")
    if st.button("Проверить подключение к CalDAV", key=f"btn_test_caldav_{block_id}", help="Проверяет доступ к календарю по YANDEX_CALENDAR_*."):
        try:
            from src.yandex_calendar_client import fetch_existing_event_keys, get_yandex_calendar_config
            cfg = get_yandex_calendar_config()
            url = (cfg.get("neuropulse_calendar_url") or cfg.get("calendar_url") or "").strip()
            if not url or not cfg.get("user") or not cfg.get("password"):
                st.error("Задайте YANDEX_CALENDAR_NEUROPULSE_URL, YANDEX_CALENDAR_USER и YANDEX_CALENDAR_APP_PASSWORD в config/.env.")
            else:
                today = datetime.utcnow().date()
                _, ok = fetch_existing_event_keys(url, today, today)
                if ok:
                    st.success("Подключение к CalDAV успешно.")
                else:
                    st.warning("Не удалось загрузить события (сеть или доступ).")
        except Exception as e:
            st.error(f"Ошибка: {e}")
    if st.button("Синхронизировать с календарём", key=f"btn_push_cal_{block_id}", help="Добавляет отсутствующие события. Дубликаты не создаются."):
        try:
            from src.yandex_calendar_client import push_grant_and_kkt_to_yandex_calendar, get_yandex_calendar_config
            cfg = get_yandex_calendar_config()
            url = (cfg.get("neuropulse_calendar_url") or cfg.get("calendar_url") or "").strip()
            if not url:
                st.error("Задайте YANDEX_CALENDAR_NEUROPULSE_URL или YANDEX_CALENDAR_URL в config/.env.")
            else:
                created, errors = push_grant_and_kkt_to_yandex_calendar(calendar_url=url, skip_existing=True)
                st.success(f"Создано событий: {created}. Ошибок: {errors}. Обновите виджет календаря или calendar.yandex.ru.")
                st.rerun()
        except Exception as e:
            st.error(f"Ошибка: {e}")


def _content_neuropulse_cal(block_id: str) -> None:
    try:
        from src.yandex_calendar_client import get_yandex_calendar_config
    except ImportError:
        st.info("Модуль Яндекс Календаря недоступен.")
        return
    _fetch_cal_failed = False
    _event_matches = lambda d, title, existing_keys: False
    cfg = get_yandex_calendar_config()
    embed_url = (cfg.get("neuropulse_embed_url") or "").strip()
    neuro_url = (cfg.get("neuropulse_calendar_url") or cfg.get("calendar_url") or "").strip()
    can_sync = bool(neuro_url and cfg.get("user") and cfg.get("password"))

    # Автосинхронизация: при первом открытии блока — двусторонняя синхронизация с календарём Нейропульс
    if _get("dashboard_neuropulse_auto_sync", True) and can_sync and not st.session_state.get("neuropulse_auto_synced"):
        try:
            from src.yandex_calendar_client import sync_neuropulse_calendar_state
            result = sync_neuropulse_calendar_state(calendar_url=neuro_url)
            if result.get("error"):
                st.session_state["neuropulse_auto_sync_failed"] = True
            else:
                st.session_state["neuropulse_auto_synced"] = True
                st.session_state["neuropulse_auto_sync_failed"] = False
                if result.get("created", 0) + result.get("pulled", 0) > 0:
                    st.rerun()
        except Exception:
            st.session_state["neuropulse_auto_sync_failed"] = True

    # Виджет календаря Яндекса (если задан embed URL).
    # Если в embed_url нет layer_ids, извлекаем ID из neuropulse_calendar_url и добавляем.
    # Заменяем week-view на month-view чтобы показывались предстоящие события.
    if embed_url:
        import re as _re
        if "layer_ids=" not in embed_url and neuro_url:
            cal_id_match = _re.search(r"/events-(\d+)/?", neuro_url)
            if cal_id_match:
                layer_id = cal_id_match.group(1)
                sep = "&" if "?" in embed_url else "?"
                embed_url = embed_url + f"{sep}layer_ids={layer_id}"
        # week → month чтобы было видно предстоящие события
        embed_url = embed_url.replace("/embed/week", "/embed/month")
        safe_url = embed_url.replace('"', "&quot;").replace("<", "").replace(">", "")
        iframe_html = (
            f'<iframe src="{safe_url}" width="800" height="600" frameborder="0" '
            'allow="microphone \'none\'; camera \'none\'" '
            'style="border: 1px solid #eee; max-width: 100%; box-sizing: border-box;"></iframe>'
        )
        st.components.v1.html(iframe_html, height=610)
    else:
        st.caption("Чтобы встроить виджет календаря, укажите **YANDEX_CALENDAR_NEUROPULSE_EMBED_URL** в config/.env (Экспорт → вставка на сайт в calendar.yandex.ru).")

    # Кнопки ручной синхронизации
    if st.session_state.get("neuropulse_auto_sync_failed"):
        st.warning("⚠️ Автосинхронизация не выполнилась. Нажмите «Синхронизировать», чтобы повторить.")
    st.caption("Двусторонняя синхронизация: изменения из календаря подтягиваются в локальные данные, локальные — в календарь.")
    try:
        from src.yandex_calendar_client import sync_neuropulse_calendar_state, load_registry, get_yandex_calendar_config
        cal_cfg = get_yandex_calendar_config()
        neuro_url = (cal_cfg.get("neuropulse_calendar_url") or cal_cfg.get("calendar_url") or "").strip()
        if neuro_url and cal_cfg.get("user") and cal_cfg.get("password"):
            registry = load_registry()
            last_sync = registry.get("last_sync")
            if last_sync:
                st.caption(f"Последняя синхронизация: {last_sync[:19].replace('T', ' ')}")
            if st.button("Синхронизировать с календарём", key="btn_sync_neuropulse_content", type="primary", help="Двусторонняя синхронизация: локальные события и календарь Нейропульс."):
                try:
                    result = sync_neuropulse_calendar_state(calendar_url=neuro_url)
                    err = result.get("error")
                    if err:
                        st.session_state["neuropulse_auto_sync_failed"] = True
                        st.error(f"Ошибка синхронизации: {err}")
                    else:
                        st.session_state["neuropulse_auto_synced"] = True
                        st.session_state["neuropulse_auto_sync_failed"] = False
                        c, u, d, p = result.get("created", 0), result.get("updated", 0), result.get("deleted", 0), result.get("pulled", 0)
                        st.success(f"Создано: {c}, обновлено: {u}, удалено: {d}, подтянуто из календаря: {p}. Обновите виджет выше или calendar.yandex.ru.")
                        st.rerun()
                except Exception as e:
                    st.session_state["neuropulse_auto_sync_failed"] = True
                    st.error(f"Ошибка синхронизации: {e}")
            with st.expander("Повторная двусторонняя синхронизация"):
                if st.button("Выполнить синхронизацию", key="btn_full_sync_neuropulse_content"):
                    try:
                        result = sync_neuropulse_calendar_state(calendar_url=neuro_url)
                        if result.get("error"):
                            st.error(f"Ошибка: {result['error']}")
                        else:
                            st.session_state["neuropulse_auto_synced"] = True
                            st.session_state["neuropulse_auto_sync_failed"] = False
                            c, u, d, p = result.get("created", 0), result.get("updated", 0), result.get("deleted", 0), result.get("pulled", 0)
                            st.success(f"Создано: {c}, обновлено: {u}, удалено: {d}, подтянуто: {p}.")
                            st.rerun()
                    except Exception as e:
                        st.error(f"Ошибка: {e}")
        else:
            st.caption("Задайте YANDEX_CALENDAR_NEUROPULSE_URL (или YANDEX_CALENDAR_URL), YANDEX_CALENDAR_USER и YANDEX_CALENDAR_APP_PASSWORD в .env для синхронизации.")
    except ImportError:
        pass

    # Список: все события гранта и ККТ (с метками Грант / ККТ).
    # force_local=True — всегда из локальных файлов, не зависит от настройки источника «Яндекс Календарь».
    st.markdown("**📅 Все события гранта и ККТ** (отмечены в календаре после синхронизации)")
    st.caption("Цвет фона: Блок 1 — голубой, Блок 2 — зелёный, Блок 3 — оранжевый.")
    days = _get("dashboard_neuropulse_days", 60)
    min_year = _get("dashboard_neuropulse_min_year", 2026)
    show_desc = _get("dashboard_neuropulse_show_desc", True)
    today = datetime.utcnow().date()
    events = _load_all_grant_and_kkt_events(days_ahead=days, min_year=min_year, force_local=True)
    if not events:
        st.caption("За выбранный период событий гранта и ККТ нет. Добавьте данные в grant_calendar.json или grant_kkt.json.")
        return

    # Проверяем доступность синхронизации для кнопок «+ Календарь»
    _add_cal_available = False
    _add_cal_neuro_url = ""
    _existing_cal_keys: set[tuple[str, str]] = set()
    try:
        from src.yandex_calendar_client import (
            add_event_to_neuropulse_calendar as _add_ev,
            get_yandex_calendar_config as _gcfg,
            event_matches_existing as _event_matches,
            fetch_existing_event_keys as _fetch_existing_keys,
            load_persisted_event_keys as _load_cached_keys,
            persist_event_keys as _persist_cached_keys,
            remember_event_key as _remember_event_key,
        )
        _existing_cal_keys = _load_cached_keys()
        _ncfg = _gcfg()
        _add_cal_neuro_url = (_ncfg.get("neuropulse_calendar_url") or _ncfg.get("calendar_url") or "").strip()
        _add_cal_available = bool(_add_cal_neuro_url and _ncfg.get("user") and _ncfg.get("password"))
        # События, уже лежащие в календаре Нейропульс — показываем ✓ вместо +
        if _add_cal_available and events:
            try:
                dates_parsed = []
                for e in events:
                    ds = (e.get("date") or "")[:10]
                    if ds:
                        dates_parsed.append(datetime.strptime(ds, "%Y-%m-%d").date())
                if dates_parsed:
                    from_date = min(dates_parsed)
                    to_date = max(dates_parsed)
                    keys, ok = _fetch_existing_keys(_add_cal_neuro_url, from_date, to_date)
                    _existing_cal_keys |= keys
                    _persist_cached_keys(_existing_cal_keys)
                    if not ok:
                        _fetch_cal_failed = True
            except (ValueError, Exception):
                _fetch_cal_failed = True
    except ImportError:
        pass
    if _fetch_cal_failed:
        st.warning("⚠️ Невозможно проверить календарь онлайн. Отметки ✓ показаны по сохранённому кэшу и могут быть не полностью актуальны.")

    _BADGE_COLORS = {
        "grant":  ("#e3f2fd", "#1565c0", "Грант"),
        "kkt":    ("#e8f5e9", "#2e7d32", "ККТ"),
        "stage":  ("#fff3e0", "#e65100", "Этап"),
        "manual": ("#f3e5f5", "#7b1fa2", "Импорт"),
    }

    for idx, ev in enumerate(events):
        d = ev.get("date", "")
        source = ev.get("source") or "grant"
        stage = ev.get("stage") or ""
        title = ev.get("title", "(без названия)")
        desc = (ev.get("description") or "").strip()
        addr = (ev.get("address") or "").strip()
        bg_color = _get_stage_bg_color(stage)
        badge_bg, badge_fg, badge_text = _BADGE_COLORS.get(source, ("#f5f5f5", "#333", source))
        days_left = (datetime.fromisoformat(d.replace("Z", "")).date() - today).days if d else 0

        border_color = bg_color if bg_color != "transparent" else "#ddd"
        row_style = (
            f"background-color:{bg_color}; padding:0.4rem 0.7rem; border-radius:8px; "
            f"margin:0.25rem 0; border-left:4px solid {border_color};"
        )
        badge_html = (
            f'<span style="display:inline-block;font-size:0.72rem;padding:0.12rem 0.4rem;'
            f'border-radius:5px;margin-right:0.4rem;font-weight:700;'
            f'background:{badge_bg};color:{badge_fg}">{badge_text}</span>'
        )
        days_html = (
            f'<span style="color:#5c6bc0;font-weight:500"> ({days_left} дн.)</span>'
            if days_left >= 0 else ""
        )
        stage_html = (
            f'<span style="font-size:0.82em;color:#666"> [{stage}]</span>' if stage else ""
        )
        header_html = (
            f'{badge_html}<strong>{d}</strong>{days_html}{stage_html}'
            f' — {title}'
        )

        ev_col, btn_col = st.columns([9, 1])
        with ev_col:
            lines_html = f'<div style="{row_style}">{header_html}</div>'
            if show_desc and desc:
                lines_html += f'<div style="font-size:0.88em;color:#555;margin-left:0.6rem;margin-bottom:0.2rem">{desc}</div>'
            if addr:
                lines_html += f'<div style="font-size:0.82em;color:#666;margin-left:0.6rem">📍 {addr}</div>'
            st.markdown(lines_html, unsafe_allow_html=True)
        with btn_col:
            btn_key = f"np_add_cal_{idx}_{d}"
            ok_key = f"np_add_cal_ok_{idx}_{d}"
            _event_date = datetime.fromisoformat(d.replace("Z", "")).date()
            _already_in_cal = not _fetch_cal_failed and _event_matches(_event_date, title or "", _existing_cal_keys)
            if st.session_state.get(ok_key) or _already_in_cal:
                st.markdown("<strong>V</strong>" if _already_in_cal else "✅", unsafe_allow_html=True)
            elif _add_cal_available:
                if st.button("＋", key=btn_key, help=f"Добавить «{title}» в Календарь Нейропульс"):
                    try:
                        _start = datetime.fromisoformat(d.replace("Z", "")).date()
                        _ok, _msg = _add_ev(
                            title=title,
                            start_date=_start,
                            description=desc,
                            address=addr,
                            save_to_json=False,
                        )
                        if _ok:
                            _remember_event_key(_start, title)
                            st.session_state[ok_key] = True
                            st.rerun()
                        else:
                            st.error(_msg)
                    except Exception as _e:
                        st.error(str(_e))


def block_neuropulse_calendar() -> None:
    _block_with_settings("📅 Календарь Нейропульс", "neuropulse_cal", _content_neuropulse_cal, _settings_neuropulse_cal)


# ---------- Блок 6: Ссылки и переходы ----------
def _settings_links(block_id: str) -> None:
    show_chat = st.checkbox("Показывать кнопку «Открыть чат»", value=_get("dashboard_links_show_chat", True), key=f"cb_chat_{block_id}")
    st.session_state["dashboard_links_show_chat"] = show_chat
    show_docs = st.checkbox("Показывать кнопку «Документация»", value=_get("dashboard_links_show_docs", True), key=f"cb_docs_{block_id}")
    st.session_state["dashboard_links_show_docs"] = show_docs
    chat_url = st.text_input("URL чата (если другой порт/хост)", value=_get("dashboard_links_chat_url", "http://localhost:8501"), key=f"url_chat_{block_id}")
    st.session_state["dashboard_links_chat_url"] = chat_url
    docs_url = st.text_input("URL документации", value=_get("dashboard_links_docs_url", ""), key=f"url_docs_{block_id}")
    st.session_state["dashboard_links_docs_url"] = docs_url


def _content_links(block_id: str) -> None:
    config = load_config()
    chat_url = _get("dashboard_links_chat_url", "http://localhost:8501")
    docs_url = _get("dashboard_links_docs_url", "")
    if _get("dashboard_links_show_chat", True):
        st.link_button("💬 Открыть чат", chat_url)
    if _get("dashboard_links_show_docs", True) and docs_url:
        st.link_button("📖 Документация", docs_url)
    elif _get("dashboard_links_show_docs", True):
        st.caption("Укажите URL документации в настройках блока.")
    disk_url = (config.get("yandex_disk") or {}).get("project_folder_url", "").strip()
    if disk_url:
        st.link_button("📁 Папка проекта на Яндекс.Диске", disk_url)


def block_links() -> None:
    _block_with_settings("🔗 Ссылки и переходы", "links", _content_links, _settings_links)


# ---------- Блок 7: Контакты по гранту ----------
def _extract_contacts_from_docx(docx_source: bytes | Path) -> list[dict]:
    """Извлекает контакты (ФИО, роль, email) из таблиц в .docx заявки на грант (команда/специалисты)."""
    import io
    from docx import Document

    if isinstance(docx_source, Path):
        doc = Document(str(docx_source))
    else:
        doc = Document(io.BytesIO(docx_source))

    # Ключевые слова для определения столбцов (нижний регистр)
    NAME_KEYS = ("фио", "ф.и.о", "имя", "фам")
    ROLE_KEYS = ("должность", "роль", "позиция", "функция", "обязанност")
    EMAIL_KEYS = ("email", "e-mail", "почта", "эл. почта", "электронная почта")

    def col_index(row_cells, keys_tuple) -> int | None:
        for i, cell in enumerate(row_cells):
            t = (cell.text or "").strip().lower()
            for k in keys_tuple:
                if k in t:
                    return i
        return None

    for table in doc.tables:
        if not table.rows:
            continue
        header = table.rows[0].cells
        texts = [(c.text or "").strip().lower() for c in header]
        i_name = col_index(header, NAME_KEYS)
        i_role = col_index(header, ROLE_KEYS)
        i_email = col_index(header, EMAIL_KEYS)
        if i_name is None and i_role is None:
            continue
        if i_name is None:
            i_name = 0
        if i_role is None:
            i_role = i_name + 1 if len(header) > i_name + 1 else i_name
        if i_email is None:
            i_email = max(i_name, i_role) + 1 if len(header) > max(i_name, i_role) + 1 else i_role

        out = []
        for row in table.rows[1:]:
            cells = row.cells
            n = len(cells)
            name = (cells[i_name].text or "").strip() if n > i_name else ""
            role = (cells[i_role].text or "").strip() if n > i_role else ""
            email = (cells[i_email].text or "").strip() if n > i_email else ""
            if name or role or email:
                out.append({"name": name or "—", "role": role or "—", "email": email or "—"})
        if out:
            return out

    return []


def _load_grant_contacts() -> list[dict]:
    """Загружает контакты из data/grant_contacts.json или из примера."""
    if CONTACTS_PATH.exists():
        try:
            with open(CONTACTS_PATH, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    if CONTACTS_EXAMPLE.exists():
        try:
            with open(CONTACTS_EXAMPLE, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []


def _settings_contacts(block_id: str) -> None:
    st.caption("Редактируемая таблица сохраняется в data/grant_contacts.json. Можно добавить столбец «Телефон», удалять строки и править ячейки.")


def _content_contacts(block_id: str) -> None:
    with st.expander("📥 Заполнить из заявки на грант (.docx)", expanded=False):
        st.caption("Загрузите файл заявки с таблицей команды (столбцы: ФИО/Имя, Должность/Роль, Email). Контакты будут записаны в data/grant_contacts.json.")
        uploaded = st.file_uploader("Файл заявки .docx", type=["docx"], key="contacts_from_application_upload")
        if uploaded and st.button("Извлечь специалистов и сохранить", key="btn_fill_contacts_from_app"):
            contacts = _extract_contacts_from_docx(uploaded.getvalue())
            if not contacts:
                st.warning("В документе не найдена таблица с командой (ожидаются столбцы типа ФИО, Должность, Email).")
            else:
                try:
                    CONTACTS_PATH.parent.mkdir(parents=True, exist_ok=True)
                    with open(CONTACTS_PATH, "w", encoding="utf-8") as f:
                        json.dump(contacts, f, ensure_ascii=False, indent=2)
                    st.success(f"Сохранено контактов: {len(contacts)}. Обновите блок или перезагрузите страницу.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Ошибка записи: {e}")

    contacts = _load_grant_contacts()
    if not contacts:
        contacts = [{"name": "", "role": "", "email": "", "chat_link": "", "phone": ""}]

    key_contacts = "dashboard_contacts_edit"
    if key_contacts not in st.session_state:
        rows = []
        for c in contacts:
            phone = (c.get("phone") or c.get("телефон") or "").strip()
            rows.append({
                "Имя": (c.get("name") or "").strip(),
                "Роль": (c.get("role") or "").strip(),
                "Email": (c.get("email") or "").strip(),
                "Макс.": (c.get("chat_link") or c.get("max_chat") or "").strip(),
                "Телефон": phone,
            })
        st.session_state[key_contacts] = rows
    else:
        for r in st.session_state[key_contacts]:
            r.setdefault("Макс.", "")
            r.setdefault("Телефон", "")

    rows = st.session_state[key_contacts]

    col_config = {
        "Имя": st.column_config.TextColumn("Имя", width="medium"),
        "Роль": st.column_config.TextColumn("Роль", width="medium"),
        "Email": st.column_config.TextColumn("Email", width="medium"),
        "Макс.": st.column_config.LinkColumn("Макс.", width="medium", help="Ссылка на чат с сотрудником (Макс)"),
        "Телефон": st.column_config.TextColumn("Телефон", width="medium"),
    }
    edited = st.data_editor(
        rows,
        width="stretch",
        hide_index=True,
        column_config=col_config,
        column_order=["Имя", "Роль", "Email", "Макс.", "Телефон"],
        key="contacts_data_editor",
        num_rows="dynamic",
    )
    st.session_state[key_contacts] = edited

    btn_col1, btn_col2 = st.columns(2)
    with btn_col1:
        if st.button("💾 Сохранить в файл", key="contacts_save", type="primary"):
            to_save = []
            for r in edited:
                rec = {
                    "name": (r.get("Имя") or "").strip(),
                    "role": (r.get("Роль") or "").strip(),
                    "email": (r.get("Email") or "").strip(),
                }
                chat_link = (r.get("Макс.") or "").strip()
                if chat_link:
                    rec["chat_link"] = chat_link
                phone = (r.get("Телефон") or "").strip()
                if phone:
                    rec["phone"] = phone
                to_save.append(rec)
            try:
                CONTACTS_PATH.parent.mkdir(parents=True, exist_ok=True)
                with open(CONTACTS_PATH, "w", encoding="utf-8") as f:
                    json.dump(to_save, f, ensure_ascii=False, indent=2)
                st.success(f"Сохранено контактов: {len(to_save)}.")
            except Exception as e:
                st.error(f"Ошибка записи: {e}")
    with btn_col2:
        if st.button("🔄 Сбросить из файла", key="contacts_reset"):
            if key_contacts in st.session_state:
                del st.session_state[key_contacts]
            st.rerun()



def block_contacts() -> None:
    _block_with_settings("👥 Контакты по гранту", "contacts", _content_contacts, _settings_contacts)


# ---------- Мониторинг гранта «НейроПульс» (этапы, бюджет, показатели) ----------
def _settings_grant_header(block_id: str) -> None:
    st.caption("Данные из data/grant_project_dashboard.json (скопируйте из grant_project_dashboard.example.json).")


def _content_grant_header(block_id: str) -> None:
    data = _load_grant_dashboard_data()
    h = data.get("header") or {}
    title = (h.get("title") or "НейроПульс").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
    date_start = (h.get("date_start") or "—").replace("<", "&lt;")
    date_end = (h.get("date_end") or "—").replace("<", "&lt;")
    stage = (h.get("current_stage") or "—").replace("<", "&lt;")
    responsible = (h.get("responsible") or "—").replace("<", "&lt;")
    updated = (h.get("updated_at") or "—").replace("<", "&lt;")
    stats_html = f"""
    <div class="np-stats-grid">
      <div class="np-stat-card">
        <span class="np-stat-label">Сроки</span>
        <span class="np-stat-value">{date_start} – {date_end}</span>
      </div>
      <div class="np-stat-card">
        <span class="np-stat-label">Этап</span>
        <span class="np-stat-value">{stage}</span>
        <span class="np-stat-badge">текущий</span>
      </div>
      <div class="np-stat-card">
        <span class="np-stat-label">Обновлено</span>
        <span class="np-stat-value">{updated}</span>
      </div>
      <div class="np-stat-card">
        <span class="np-stat-label">Ответственный</span>
        <span class="np-stat-value">{responsible}</span>
      </div>
    </div>
    <p class="np-grant-title">{title}</p>
    <style>
    .np-stats-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 16px; margin-bottom: 12px; }}
    .np-stat-card {{ background: #fff; border-radius: 20px; padding: 16px 20px; display: flex; flex-direction: column; border: 1px solid #e2e8f0; box-shadow: 0 4px 12px rgba(0,0,0,0.02); }}
    .np-stat-label {{ font-size: 14px; color: #64748b; margin-bottom: 4px; }}
    .np-stat-value {{ font-size: 16px; font-weight: 600; color: #0f172a; }}
    .np-stat-badge {{ font-size: 12px; background: #e0f2fe; color: #0369a1; padding: 2px 10px; border-radius: 30px; align-self: flex-start; margin-top: 8px; }}
    .np-grant-title {{ font-weight: 600; color: #0f172a; margin: 0; font-size: 1.05rem; }}
    </style>
    """
    st.markdown(stats_html, unsafe_allow_html=True)


def block_grant_header() -> None:
    _block_with_settings("📌 Общая информация по гранту", "grant_header", _content_grant_header, _settings_grant_header)


def _settings_stages(block_id: str) -> None:
    st.caption("Таблица этапов в data/grant_project_dashboard.json → stages. Статусы: not_started, in_progress, done, overdue.")


def _content_stages(block_id: str) -> None:
    data = _load_grant_dashboard_data()
    stages = data.get("stages") or []
    if not stages:
        st.info("Добавьте данные в data/grant_project_dashboard.json (секция stages). Пример: data/grant_project_dashboard.example.json")
        return

    def _esc(s: str) -> str:
        return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    table_rows = []
    for s in stages:
        status_key = (s.get("status") or "not_started").replace(" ", "_")
        status_label = STATUS_LABELS.get(s.get("status") or "not_started", s.get("status") or "not_started")
        pct = min(100, max(0, int(s.get("percent", 0))))
        stage_name = _esc(s.get("stage", ""))
        activity = _esc(s.get("activity", ""))
        plan_dates = _esc(f"{s.get('plan_start') or '—'} — {s.get('plan_end') or '—'}")
        fact_dates = _esc(f"{s.get('fact_start') or '—'} — {s.get('fact_end') or '—'}")
        table_rows.append(
            f'<tr><td>{stage_name}</td><td>{activity}</td><td>{plan_dates}</td><td>{fact_dates}</td>'
            f'<td><span class="np-status-dot np-{status_key}"></span>{_esc(status_label)}</td>'
            f'<td class="np-progress-cell"><div class="np-progress-bar-inline"><div class="np-progress-fill-inline" style="width:{pct}%"></div></div></td></tr>'
        )
    html = (
        '<table class="np-stages-table"><thead><tr>'
        '<th>Этап</th><th>Мероприятие</th><th>План: даты</th><th>Факт: даты</th><th>Статус</th><th>%</th>'
        '</tr></thead><tbody>' + "".join(table_rows) + "</tbody></table>"
    )
    st.markdown(html, unsafe_allow_html=True)


def block_stages() -> None:
    _block_with_settings("📊 Прогресс по этапам", "grant_stages", _content_stages, _settings_stages)


def _settings_budget(block_id: str) -> None:
    st.caption("Бюджет в data/grant_project_dashboard.json → budget. Обновляйте fact в items.")


def _content_budget(block_id: str) -> None:
    data = _load_grant_dashboard_data()
    budget = data.get("budget") or {}
    total_plan = budget.get("total_plan", 0)
    total_fact = budget.get("total_fact", 0)
    items = budget.get("items") or []
    if not items and not total_plan:
        st.info("Добавьте бюджет в data/grant_project_dashboard.json (секция budget).")
        return
    pct = (total_fact / total_plan * 100) if total_plan else 0
    pct_clamped = min(100, pct)
    plan_fmt = f"{total_plan:,.0f}".replace(",", " ")
    fact_fmt = f"{total_fact:,.0f}".replace(",", " ")
    st.markdown(
        f'<div class="np-budget-overview">'
        f'<div class="np-budget-total"><span>Освоено: {fact_fmt} / {plan_fmt} руб.</span>'
        f'<div class="np-progress-bar"><div class="np-progress-fill" style="width:{pct_clamped}%"></div></div>'
        f'<span class="np-budget-pct">Освоение: {pct:.1f}%</span></div></div>'
        '<style>.np-budget-overview { margin-bottom: 12px; } '
        '.np-budget-total { display: flex; flex-direction: column; gap: 8px; } '
        '.np-budget-total > span:first-child { font-weight: 500; color: #0f172a; } '
        '.np-progress-bar { width: 100%; height: 8px; background: #e2e8f0; border-radius: 10px; overflow: hidden; } '
        '.np-progress-fill { height: 100%; background: #3b82f6; border-radius: 10px; transition: width 0.3s; } '
        '.np-budget-pct { font-size: 14px; color: #64748b; }</style>',
        unsafe_allow_html=True,
    )
    rows = []
    for it in items:
        plan = it.get("plan", 0)
        fact = it.get("fact", 0)
        rows.append({"Статья": it.get("name", ""), "План (руб.)": plan, "Факт (руб.)": fact, "% освоения": (fact / plan * 100) if plan else 0})
        for ch in it.get("children") or []:
            rows.append({"Статья": "  • " + (ch.get("name") or ""), "План (руб.)": ch.get("plan", 0), "Факт (руб.)": ch.get("fact", 0), "% освоения": (ch.get("fact", 0) / ch.get("plan", 1) * 100) if ch.get("plan") else 0})
    st.dataframe(rows, width="stretch", hide_index=True)


def block_budget() -> None:
    _block_with_settings("💰 Финансовый мониторинг", "grant_budget", _content_budget, _settings_budget)


def _settings_indicators(block_id: str) -> None:
    st.caption("Показатели в data/grant_project_dashboard.json → indicators.")


def _content_indicators(block_id: str) -> None:
    data = _load_grant_dashboard_data()
    indicators = data.get("indicators") or []
    if not indicators:
        st.info("Добавьте показатели в data/grant_project_dashboard.json → indicators.")
        return
    # Первые 3 показателя — крупные KPI-карточки
    kpi_items = indicators[:3]
    kpi_html_parts = []
    for ind in kpi_items:
        target = ind.get("target", 0)
        fact = ind.get("fact", 0)
        suffix = (ind.get("suffix") or "")
        name = (ind.get("name") or "").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")
        kpi_html_parts.append(
            f'<div class="np-kpi-card"><div class="np-kpi-value">{fact}{suffix}</div>'
            f'<div class="np-kpi-label">{name}</div><div class="np-kpi-target">цель {target}{suffix}</div></div>'
        )
    if kpi_html_parts:
        st.markdown(
            '<div class="np-kpi-grid">' + "".join(kpi_html_parts) + '</div>'
            '<style>.np-kpi-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(160px, 1fr)); gap: 16px; margin-bottom: 16px; } '
            '.np-kpi-card { background: #fff; border-radius: 20px; padding: 24px; text-align: center; border: 1px solid #e2e8f0; box-shadow: 0 4px 12px rgba(0,0,0,0.02); } '
            '.np-kpi-value { font-size: 32px; font-weight: 700; color: #0f172a; line-height: 1; } '
            '.np-kpi-label { font-size: 14px; color: #64748b; margin: 8px 0 4px; } '
            '.np-kpi-target { font-size: 13px; color: #94a3b8; }</style>',
            unsafe_allow_html=True,
        )
    rows = []
    for ind in indicators:
        target = ind.get("target", 0)
        fact = ind.get("fact", 0)
        suffix = ind.get("suffix", "")
        pct = (fact / target * 100) if target else 0
        rows.append({
            "Показатель": ind.get("name", ""),
            "Цель": f"{target}{suffix}",
            "Факт": f"{fact}{suffix}",
            "% выполнения": round(pct, 1),
        })
    st.dataframe(rows, width="stretch", hide_index=True)


def block_indicators() -> None:
    _block_with_settings("👥 Результаты с участниками", "grant_indicators", _content_indicators, _settings_indicators)


def _settings_info_activity(block_id: str) -> None:
    st.caption("Данные в data/grant_project_dashboard.json → info_activity.")


def _content_info_activity(block_id: str) -> None:
    data = _load_grant_dashboard_data()
    ia = data.get("info_activity") or {}
    if not ia:
        st.info("Добавьте секцию info_activity в data/grant_project_dashboard.json.")
        return
    reach_t = ia.get("reach_target", 0) or 1
    reach_f = ia.get("reach_fact", 0)
    neuro_t = ia.get("neurofest_target", 100) or 1
    neuro_f = ia.get("neurofest_fact", 0)
    pub = ia.get("publications", 0)
    round_n = ia.get("round_table_participants", 0)
    reach_pct = min(100, round(100 * reach_f / reach_t)) if reach_t else 0
    neuro_pct = min(100, round(100 * neuro_f / neuro_t)) if neuro_t else 0
    reach_deg = round(3.6 * reach_pct)
    neuro_deg = round(3.6 * neuro_pct)

    def _circle_card(deg: float, center_text: str, label: str, target_text: str) -> str:
        style = f"background: conic-gradient(#3b82f6 0deg {deg}deg, #e2e8f0 {deg}deg 360deg);"
        return (
            f'<div class="np-circle-card">'
            f'<div class="np-circle-ring" style="{style}"><span class="np-circle-inner">{center_text}</span></div>'
            f'<div class="np-circle-label">{label}</div>'
            f'<div class="np-circle-target">{target_text}</div></div>'
        )

    cards_html = [
        _circle_card(reach_deg, f"{reach_f:,}".replace(",", " "), "Охват кампании (ЯНАО)", f"цель {reach_t:,}".replace(",", " ")),
        _circle_card(0, str(pub), "Публикации в СМИ", ""),
        _circle_card(0, str(round_n), "Участники круглого стола", ""),
        _circle_card(neuro_deg, str(neuro_f), "НейроФест: участники", f"план {neuro_t}"),
    ]
    st.markdown('<div class="np-circle-grid">' + "".join(cards_html) + "</div>", unsafe_allow_html=True)


def block_info_activity() -> None:
    _block_with_settings("📢 Информационная активность", "grant_info_activity", _content_info_activity, _settings_info_activity)


# ---------- Паспорт: паспорт проекта и краткая история диалогов (автосохранение каждые 15 мин) ----------
def _write_passport_project() -> None:
    """Формирует и записывает паспорт проекта в Паспорт/паспорт_проекта.md."""
    PASSPORT_DIR.mkdir(parents=True, exist_ok=True)
    data = _load_grant_dashboard_data()
    lines = ["# Паспорт проекта\n", f"*Обновлено: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n"]
    h = (data.get("header") or {})
    lines.append("## Шапка\n")
    lines.append(f"- **Название:** {h.get('title', '—')}\n")
    lines.append(f"- **Период:** {h.get('date_start', '—')} — {h.get('date_end', '—')} ({h.get('months_duration', '—')} мес.)\n")
    lines.append(f"- **Текущий этап:** {h.get('current_stage', '—')}\n")
    lines.append(f"- **Ответственный:** {h.get('responsible', '—')}\n")
    stages = data.get("stages") or []
    lines.append("\n## Этапы\n")
    for s in stages:
        st_label = STATUS_LABELS.get(s.get("status"), s.get("status", ""))
        lines.append(f"- **{s.get('stage', '')}** — {s.get('activity', '')} | {s.get('plan_start', '')}–{s.get('plan_end', '')} | {st_label}\n")
    budget = data.get("budget") or {}
    lines.append("\n## Бюджет\n")
    lines.append(f"- План всего: {budget.get('total_plan', 0):,}\n")
    lines.append(f"- Факт: {budget.get('total_fact', 0):,}\n")
    lines.append(f"- Грант: {budget.get('grant_amount', 0):,} | Собственный вклад: {budget.get('own_contribution', 0):,}\n")
    ind = data.get("indicators") or []
    lines.append("\n## Показатели\n")
    for i in ind:
        lines.append(f"- {i.get('name', '')}: цель {i.get('target', '—')}{i.get('suffix', '')} | факт {i.get('fact', 0)}\n")
    info = data.get("info_activity") or {}
    lines.append("\n## Информационная активность\n")
    lines.append(f"- Охват: цель {info.get('reach_target', 0)}, факт {info.get('reach_fact', 0)}\n")
    lines.append(f"- Публикации: {info.get('publications', 0)} | Круглый стол: {info.get('round_table_participants', 0)}\n")
    lines.append(f"- НейроФест: цель {info.get('neurofest_target', 0)}, факт {info.get('neurofest_fact', 0)}\n")
    try:
        PASSPORT_PROJECT_FILE.write_text("".join(lines), encoding="utf-8")
    except Exception as e:
        logger.warning("Ошибка записи паспорта проекта: %s", e)


def _write_dialogue_history(session_messages: list[dict]) -> None:
    """Формирует и записывает краткую историю диалогов в Паспорт/история_диалогов.md."""
    PASSPORT_DIR.mkdir(parents=True, exist_ok=True)
    lines = ["# Краткая история диалогов\n", f"*Обновлено: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n"]
    # Сводка по аудит-логу
    lines.append("## Аудит запросов (data/chat_audit.jsonl)\n")
    if not AUDIT_LOG.exists():
        lines.append("Записей пока нет.\n")
    else:
        records = []
        try:
            with open(AUDIT_LOG, encoding="utf-8") as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        records.append(json.loads(line))
                    except json.JSONDecodeError:
                        continue
        except Exception as e:
            lines.append(f"Ошибка чтения: {e}\n")
        if not records:
            lines.append("Записей нет.\n")
        else:
            by_day = {}
            for r in records[-500:]:
                ts = (r.get("ts") or "")[:10]
                by_day.setdefault(ts, {"count": 0, "prompt_len": 0, "reply_len": 0})
                by_day[ts]["count"] += 1
                by_day[ts]["prompt_len"] += r.get("prompt_len", 0)
                by_day[ts]["reply_len"] += r.get("reply_len", 0)
            lines.append("| Дата | Запросов | Символов запрос | Символов ответ |\n")
            lines.append("| --- | --- | --- | --- |\n")
            for day in sorted(by_day.keys(), reverse=True)[:30]:
                v = by_day[day]
                lines.append(f"| {day} | {v['count']} | {v['prompt_len']} | {v['reply_len']} |\n")
    # Текущая сессия дашборда (кратко)
    lines.append("\n## Текущая сессия дашборда (снимок)\n")
    if not session_messages:
        lines.append("Сообщений в текущей сессии нет.\n")
    else:
        for i, msg in enumerate(session_messages[-50:], 1):
            role = msg.get("role", "")
            content = (msg.get("content") or "")[:300]
            if len((msg.get("content") or "")) > 300:
                content += "..."
            content = content.replace("\n", " ")
            lines.append(f"- **{i}. {role}:** {content}\n")
    try:
        PASSPORT_HISTORY_FILE.write_text("".join(lines), encoding="utf-8")
    except Exception as e:
        logger.warning("Ошибка записи истории диалогов: %s", e)


def _passport_save_if_due(session_messages: list) -> None:
    """Если прошло 15 минут с последнего сохранения — записывает паспорт и историю, обновляет метку времени."""
    PASSPORT_DIR.mkdir(parents=True, exist_ok=True)
    now = datetime.now()
    last_save = None
    if PASSPORT_LAST_SAVE_FILE.exists():
        try:
            last_save = datetime.fromisoformat(PASSPORT_LAST_SAVE_FILE.read_text(encoding="utf-8").strip())
        except Exception:
            pass
    if last_save is None or (now - last_save).total_seconds() >= PASSPORT_INTERVAL_MINUTES * 60:
        _write_passport_project()
        _write_dialogue_history(session_messages)
        try:
            PASSPORT_LAST_SAVE_FILE.write_text(now.isoformat(), encoding="utf-8")
        except Exception as e:
            logger.warning("Ошибка записи метки автосохранения Паспорт: %s", e)
        git_push_if_changes(commit_message="Автосохранение: паспорт и история диалогов")


# ---------- Проверка отчёта агентом ----------
def _extract_text_from_uploaded_file(uploaded_file) -> str | None:
    """Извлекает текст из загруженного файла (.txt или .docx). Макс. 50_000 символов."""
    max_chars = 50_000
    try:
        if uploaded_file.name.lower().endswith(".txt"):
            text = uploaded_file.getvalue().decode("utf-8", errors="replace")
        elif uploaded_file.name.lower().endswith(".docx"):
            import io
            from docx import Document
            doc = Document(io.BytesIO(uploaded_file.getvalue()))
            text = "\n".join(p.text for p in doc.paragraphs)
        else:
            return None
        return text[:max_chars] + ("…" if len(text) > max_chars else "")
    except Exception as e:
        logger.warning("Ошибка извлечения текста из файла: %s", e)
        return None


def _block_check_report() -> None:
    """Кнопка «Проверить отчёт агентом»: загрузка файла, отправка текста агенту."""
    if not _ensure_conversation():
        return
    with st.expander("📄 Проверить отчёт агентом", expanded=False):
        st.caption("Загрузите .txt или .docx — агент проверит отчёт на полноту и соответствие грантовым требованиям.")
        uploaded = st.file_uploader("Файл отчёта", type=["txt", "docx"], key="dashboard_report_upload")
        if uploaded and st.button("Отправить на проверку", key="btn_check_report"):
            text = _extract_text_from_uploaded_file(uploaded)
            if not text:
                st.error("Не удалось прочитать текст из файла. Поддерживаются .txt и .docx.")
                return
            prompt = "Проверь этот отчёт на полноту и соответствие грантовым требованиям. Укажи, если чего-то не хватает или есть замечания.\n\n---\n\n" + text
            st.session_state.dashboard_messages.append({"role": "user", "content": f"[Проверка отчёта: {uploaded.name}]"})
            with st.spinner("Агент проверяет отчёт..."):
                reply = _send_to_agent(prompt)
                st.session_state.dashboard_messages.append({"role": "assistant", "content": reply or "Ошибка запроса к агенту."})
            st.success("Ответ агента добавлен в диалог ниже.")
            st.rerun()


# ---------- Окно диалога с агентом ----------
def block_chat() -> None:
    """Окно диалога с агентом. При нажатии на быстрые действия по гранту их формулировки (prompt из config) отправляются агенту и ответ показывается здесь."""
    st.subheader("💬 Диалог с агентом")
    if not _ensure_conversation():
        st.warning("Настройте API (api_key, folder_id в config или .env) для диалога с агентом.")
        return

    if "dashboard_messages" not in st.session_state:
        st.session_state.dashboard_messages = []

    _block_check_report()

    # Обработка очереди от быстрых действий: формулировки (prompt) из кнопок уходят сюда
    if "dashboard_prompt_to_send" in st.session_state and st.session_state.dashboard_prompt_to_send:
        prompt = (st.session_state.dashboard_prompt_to_send.pop(0) or "").strip()
        if prompt:
            st.session_state.dashboard_messages.append({"role": "user", "content": prompt})
            with st.spinner("Ответ агента..."):
                reply = _send_to_agent(prompt)
                st.session_state.dashboard_messages.append({"role": "assistant", "content": reply or "Ошибка запроса. Проверьте настройки API в config/.env."})
        st.rerun()

    for msg in st.session_state.dashboard_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if prompt := st.chat_input("Введите сообщение или нажмите быстрый запрос справа"):
        st.session_state.dashboard_messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        with st.chat_message("assistant"):
            with st.spinner("Думаю..."):
                reply = _send_to_agent(prompt)
                text = reply if reply else "Ошибка запроса к агенту (пустой ответ). " + _agent_config_hint()
                st.markdown(text)
                st.session_state.dashboard_messages.append({"role": "assistant", "content": text})
        st.rerun()

    with st.expander("📅 Добавить событие в календарь Нейропульс", expanded=False):
        try:
            from src.yandex_calendar_client import add_event_to_neuropulse_calendar, get_yandex_calendar_config
            cfg = get_yandex_calendar_config()
            if (cfg.get("neuropulse_calendar_url") or cfg.get("calendar_url")) and cfg.get("user") and cfg.get("password"):
                ev_date = st.date_input("Дата", value=datetime.utcnow().date(), key="add_cal_ev_date")
                ev_title = st.text_input("Название", placeholder="Например: Встреча с партнёром", key="add_cal_ev_title")
                ev_desc = st.text_area("Описание", placeholder="Опционально", key="add_cal_ev_desc", height=60)
                ev_addr = st.text_input("Место / адрес", placeholder="Опционально", key="add_cal_ev_addr")
                save_to_json = st.checkbox("Добавить в grant_calendar.json", value=True, key="add_cal_ev_save_json")
                if st.button("Создать событие в календаре", key="add_cal_ev_btn"):
                    if not (ev_title or "").strip():
                        st.warning("Укажите название события.")
                    else:
                        ok, msg = add_event_to_neuropulse_calendar(
                            title=(ev_title or "").strip(),
                            start_date=ev_date,
                            description=(ev_desc or "").strip(),
                            address=(ev_addr or "").strip(),
                            save_to_json=save_to_json,
                        )
                        if ok:
                            st.success(msg)
                        else:
                            st.error(msg)
            else:
                st.caption("Задайте YANDEX_CALENDAR_NEUROPULSE_URL (или YANDEX_CALENDAR_URL), YANDEX_CALENDAR_USER и YANDEX_CALENDAR_APP_PASSWORD в .env.")
        except ImportError:
            st.caption("Модуль Яндекс Календаря недоступен.")


def block_chat_and_tools() -> None:
    """Диалог с агентом и быстрые действия в одной колонке (2:1)."""
    chat_col, tools_col = st.columns([2, 1])
    with chat_col:
        block_chat()
    with tools_col:
        block_tools()


# Реестр блоков дашборда для настраиваемого расположения (перенос между колонками и порядок)
DEFAULT_LAYOUT_LEFT = [
    "grant_header", "stages", "kkt", "budget", "indicators", "risks", "info_activity",
    "schedule", "reminders", "audit", "agent_stats", "status",
]
DEFAULT_LAYOUT_RIGHT = [
    "neuropulse_calendar", "chat_and_tools", "contacts", "notifications",
    "communications", "documents_archive", "links", "vector_store", "parser",
]


def _get_block_registry() -> dict:
    """Возвращает реестр блоков: id -> (название, функция отрисовки)."""
    return {
        "grant_header": ("Шапка гранта", block_grant_header),
        "stages": ("Этапы", block_stages),
        "kkt": ("Ключевые контрольные точки", block_kkt),
        "budget": ("Бюджет", block_budget),
        "indicators": ("Показатели", block_indicators),
        "risks": ("Риски", block_risks),
        "info_activity": ("Информационная активность", block_info_activity),
        "schedule": ("Ближайшие сроки", block_schedule),
        "calendar_month": ("Календарь (месяц)", block_calendar_month),
        "reminders": ("Напоминания", block_reminders),
        "neuropulse_calendar": ("Календарь Нейропульс", block_neuropulse_calendar),
        "audit": ("Аудит чата", block_audit),
        "agent_stats": ("Статистика агента", block_agent_stats),
        "status": ("Настройки", block_status),
        "chat_and_tools": ("Диалог с агентом и быстрые действия", block_chat_and_tools),
        "contacts": ("Контакты по гранту", block_contacts),
        "notifications": ("Настройки уведомлений", block_notifications),
        "communications": ("Коммуникации с грантодателем", block_communications),
        "documents_archive": ("Файловый архив", block_documents_archive),
        "links": ("Ссылки", block_links),
        "vector_store": ("Vector Store", block_vector_store),
        "parser": ("🔀 Агент Парсер", block_parser),
    }


def _load_dashboard_layout() -> tuple[list[str], list[str]]:
    """Загружает раскладку из data/dashboard_layout.json или возвращает умолчание."""
    registry = _get_block_registry()
    valid_ids = set(registry)
    if DASHBOARD_LAYOUT_PATH.exists():
        try:
            with open(DASHBOARD_LAYOUT_PATH, encoding="utf-8") as f:
                data = json.load(f)
            left = [x for x in (data.get("left") or []) if x in valid_ids]
            right = [x for x in (data.get("right") or []) if x in valid_ids]
            if left or right:
                used = set(left) | set(right)
                for bid in valid_ids:
                    if bid not in used:
                        right.append(bid)
                return (left, right)
        except Exception:
            pass
    return (list(DEFAULT_LAYOUT_LEFT), list(DEFAULT_LAYOUT_RIGHT))


def _save_dashboard_layout(left: list[str], right: list[str]) -> None:
    """Сохраняет раскладку в data/dashboard_layout.json (сохраняет поле hidden)."""
    DASHBOARD_LAYOUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    existing: dict = {}
    if DASHBOARD_LAYOUT_PATH.exists():
        try:
            with open(DASHBOARD_LAYOUT_PATH, encoding="utf-8") as f:
                existing = json.load(f)
        except Exception:
            pass
    existing["left"] = left
    existing["right"] = right
    with open(DASHBOARD_LAYOUT_PATH, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)


def _load_hidden_blocks() -> set[str]:
    """Загружает множество скрытых блоков из dashboard_layout.json."""
    if DASHBOARD_LAYOUT_PATH.exists():
        try:
            with open(DASHBOARD_LAYOUT_PATH, encoding="utf-8") as f:
                data = json.load(f)
            return set(data.get("hidden") or [])
        except Exception:
            pass
    return set()


def _save_hidden_blocks(hidden: set[str]) -> None:
    """Сохраняет список скрытых блоков в dashboard_layout.json."""
    DASHBOARD_LAYOUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    existing: dict = {}
    if DASHBOARD_LAYOUT_PATH.exists():
        try:
            with open(DASHBOARD_LAYOUT_PATH, encoding="utf-8") as f:
                existing = json.load(f)
        except Exception:
            pass
    existing["hidden"] = sorted(hidden)
    with open(DASHBOARD_LAYOUT_PATH, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)


def _is_block_hidden(block_id: str) -> bool:
    """Проверяет, скрыт ли блок (берёт из session_state для скорости)."""
    return block_id in st.session_state.get("dashboard_hidden_blocks", set())


def _render_layout_editor() -> None:
    """Рендерит экспандер «Расположение блоков»: перенос между колонками и изменение порядка (аналог drag-and-drop)."""
    registry = _get_block_registry()
    key_layout = "dashboard_layout_blocks"
    st.session_state[key_layout] = _load_dashboard_layout()
    left_ids, right_ids = st.session_state[key_layout]

    hidden_set = st.session_state.get("dashboard_hidden_blocks", _load_hidden_blocks())

    with st.expander("📐 Расположение блоков (перенос, порядок, видимость)", expanded=False):
        st.caption("Меняйте порядок блоков и колонку (левая/правая). Скрытые блоки помечены 🙈 — нажмите 👁 чтобы вернуть на дашборд.")
        lcol, rcol = st.columns(2)
        with lcol:
            st.markdown("**Левая колонка**")
            for i, bid in enumerate(left_ids):
                if bid not in registry:
                    continue
                title = registry[bid][0]
                is_hid = bid in hidden_set
                label = f"🙈 ~~{title}~~" if is_hid else f"• {title}"
                # 4 колонки вместо 5: название | ↑↓ | → | 👁
                name_col, updown_col, move_col, hide_col = st.columns([4, 1, 1, 1])
                with name_col:
                    st.markdown(label)
                with updown_col:
                    if i > 0 and st.button("↑", key=f"left_up_{bid}", help="Вверх"):
                        left_ids[i], left_ids[i - 1] = left_ids[i - 1], left_ids[i]
                        _save_dashboard_layout(left_ids, right_ids)
                        st.rerun()
                    if i < len(left_ids) - 1 and st.button("↓", key=f"left_down_{bid}", help="Вниз"):
                        left_ids[i], left_ids[i + 1] = left_ids[i + 1], left_ids[i]
                        _save_dashboard_layout(left_ids, right_ids)
                        st.rerun()
                with move_col:
                    if st.button("→", key=f"left_to_right_{bid}", help="В правую колонку"):
                        left_ids.remove(bid)
                        right_ids.append(bid)
                        _save_dashboard_layout(left_ids, right_ids)
                        st.rerun()
                with hide_col:
                    eye_label = "👁" if is_hid else "🙈"
                    eye_help = "Показать блок" if is_hid else "Скрыть блок"
                    if st.button(eye_label, key=f"left_hide_{bid}", help=eye_help):
                        if is_hid:
                            hidden_set.discard(bid)
                        else:
                            hidden_set.add(bid)
                        _save_hidden_blocks(hidden_set)
                        st.session_state["dashboard_hidden_blocks"] = hidden_set
                        st.rerun()
        with rcol:
            st.markdown("**Правая колонка**")
            for i, bid in enumerate(right_ids):
                if bid not in registry:
                    continue
                title = registry[bid][0]
                is_hid = bid in hidden_set
                label = f"🙈 ~~{title}~~" if is_hid else f"• {title}"
                # 4 колонки вместо 5: название | ← | ↑↓ | 👁
                name_col, move_col, updown_col, hide_col = st.columns([4, 1, 1, 1])
                with name_col:
                    st.markdown(label)
                with move_col:
                    if st.button("←", key=f"right_to_left_{bid}", help="В левую колонку"):
                        right_ids.remove(bid)
                        left_ids.append(bid)
                        _save_dashboard_layout(left_ids, right_ids)
                        st.rerun()
                with updown_col:
                    if i > 0 and st.button("↑", key=f"right_up_{bid}", help="Вверх"):
                        right_ids[i], right_ids[i - 1] = right_ids[i - 1], right_ids[i]
                        _save_dashboard_layout(left_ids, right_ids)
                        st.rerun()
                    if i < len(right_ids) - 1 and st.button("↓", key=f"right_down_{bid}", help="Вниз"):
                        right_ids[i], right_ids[i + 1] = right_ids[i + 1], right_ids[i]
                        _save_dashboard_layout(left_ids, right_ids)
                        st.rerun()
                with hide_col:
                    eye_label = "👁" if is_hid else "🙈"
                    eye_help = "Показать блок" if is_hid else "Скрыть блок"
                    if st.button(eye_label, key=f"right_hide_{bid}", help=eye_help):
                        if is_hid:
                            hidden_set.discard(bid)
                        else:
                            hidden_set.add(bid)
                        _save_hidden_blocks(hidden_set)
                        st.session_state["dashboard_hidden_blocks"] = hidden_set
                        st.rerun()
        if st.button("Сбросить к умолчанию", key="layout_reset"):
            st.session_state[key_layout] = (list(DEFAULT_LAYOUT_LEFT), list(DEFAULT_LAYOUT_RIGHT))
            _save_dashboard_layout(DEFAULT_LAYOUT_LEFT, DEFAULT_LAYOUT_RIGHT)
            st.rerun()


def _inject_theme_css() -> None:
    """Подключает тему НейроПульс: neuropulse_theme.css (рядом с app.py) + шрифт Inter."""
    _dashboard_dir = Path(__file__).resolve().parent
    theme_path = _dashboard_dir / "neuropulse_theme.css"
    if theme_path.exists():
        try:
            css = theme_path.read_text(encoding="utf-8")
            st.markdown(f"<style id=\"neuropulse-theme\">\n{css}\n</style>", unsafe_allow_html=True)
        except Exception as e:
            logger.warning("Не удалось загрузить тему дашборда %s: %s", theme_path, e)
            _inject_theme_fallback()
    else:
        logger.warning("Файл темы не найден: %s", theme_path)
        _inject_theme_fallback()


def _inject_theme_fallback() -> None:
    """Минимальные стили и шрифт Inter, если neuropulse_theme.css недоступен."""
    st.markdown(
        "<link rel=\"preconnect\" href=\"https://fonts.googleapis.com\">"
        "<link rel=\"preconnect\" href=\"https://fonts.gstatic.com\" crossorigin>"
        "<link href=\"https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap\" rel=\"stylesheet\">"
        "<style>body,.stApp{background:#f8fafc!important;} .main .block-container{font-family:'Inter',sans-serif!important;}</style>",
        unsafe_allow_html=True,
    )


def main() -> None:
    st.set_page_config(
        page_title="Дашборд — Грантовый контролёр НейроПульс",
        page_icon="📋",
        layout="wide",
    )
    _inject_theme_css()
    _inject_block_palette_css()
    # Автосохранение Паспорт каждые 15 минут
    _passport_save_if_due(st.session_state.get("dashboard_messages", []))
    # Подтверждение после перезапуска (диалог с агентом не трогаем — он в session_state и сохраняется при rerun)
    if st.session_state.get("dashboard_restart_requested"):
        n_msg = len(st.session_state.get("dashboard_messages", []))
        st.success("✅ Дашборд перезапущен. Данные обновлены." + (f" Диалог с агентом сохранён ({n_msg} сообщ.)." if n_msg else ""))
        st.session_state["dashboard_restart_requested"] = False
    head_col1, head_col2 = st.columns([4, 1])
    with head_col1:
        st.title("📋 Дашборд проекта НейроПульс")
        st.caption("Мониторинг гранта, сроки, диалог с агентом. Расположение блоков настраивается в экспандере «Расположение блоков». Дашборд не использует микрофон — не мешает голосовым звонкам в других вкладках (Chrome, Google Meet).")
    with head_col2:
        if st.button("🔄 Перезапустите дашборд", type="primary", help="Обновить данные и перезагрузить страницу. Диалог с агентом сохраняется."):
            st.session_state["dashboard_restart_requested"] = True
            st.rerun()

    left_ids, right_ids = _load_dashboard_layout()
    # Загружаем список скрытых блоков в session_state (один раз за rerun)
    if "dashboard_hidden_blocks" not in st.session_state:
        st.session_state["dashboard_hidden_blocks"] = _load_hidden_blocks()
    hidden_blocks = st.session_state["dashboard_hidden_blocks"]
    st.session_state["dashboard_layout_left"] = left_ids
    st.session_state["dashboard_layout_right"] = right_ids

    registry = _get_block_registry()

    # Блок «Общая информация по гранту» — на всю ширину (1 строка = ширина 2 колонок)
    if "grant_header" in registry and "grant_header" not in hidden_blocks:
        if "grant_header" in left_ids or "grant_header" in right_ids:
            st.session_state["dashboard_current_block"] = {"bid": "grant_header", "side": "full", "index": 0, "total": 1}
            registry["grant_header"][1]()
            st.divider()

    # Маркер для CSS (display:none — в DOM для :has(), но без влияния на layout)
    st.markdown(
        '<div class="np-dashboard-columns-marker" aria-hidden="true" style="display:none;"></div>',
        unsafe_allow_html=True,
    )
    col1, col2 = st.columns(2)

    # В колонках не показываем grant_header — он уже отрисован выше на всю ширину
    left_ids = [b for b in left_ids if b != "grant_header"]
    right_ids = [b for b in right_ids if b != "grant_header"]

    with col1:
        for i, bid in enumerate(left_ids):
            if bid not in registry:
                continue
            if bid in hidden_blocks:
                continue
            st.session_state["dashboard_current_block"] = {"bid": bid, "side": "left", "index": i, "total": len(left_ids)}
            registry[bid][1]()
            if i < len(left_ids) - 1:
                st.divider()

    with col2:
        for i, bid in enumerate(right_ids):
            if bid not in registry:
                continue
            if bid in hidden_blocks:
                continue
            st.session_state["dashboard_current_block"] = {"bid": bid, "side": "right", "index": i, "total": len(right_ids)}
            registry[bid][1]()
            if i < len(right_ids) - 1:
                st.divider()

    # Экспандер «Расположение блоков» — внизу, чтобы не создавать скрытые элементы между grant_header и колонками
    _render_layout_editor()


if __name__ == "__main__":
    _register_atexit()
    main()
