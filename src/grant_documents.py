"""
Генерация документов по событиям календаря гранта.
Запрос к агенту → текст → Word (.docx) → сохранение, отправка на почту, загрузка в Vector Store.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from src.agent_api_client import PROJECT_ROOT, load_config
from src.yandex_ai_client import ask, get_yandex_ai_config

logger = logging.getLogger(__name__)

GENERATED_DIR = PROJECT_ROOT / "data" / "generated_docs"
DOCUMENT_TYPES = ("report", "event_scenario", "algorithm")


def _doc_type_prompt(doc_type: str, event_title: str, event_description: str) -> str:
    if doc_type == "report":
        kind = "требуемый отчётный документ"
    elif doc_type == "event_scenario":
        kind = "сценарий проводимого мероприятия"
    elif doc_type == "algorithm":
        kind = "алгоритм действий или структура (пошаговый план)"
    else:
        kind = "документ (структура и содержание)"
    return (
        f"Для события гранта «{event_title}»"
        + (f". Описание: {event_description}." if event_description else ".")
        + f" Сформируй {kind} по гранту губернатора и проекту Нейропульс. "
        "Дай полный текст, готовый для вставки в документ Word: заголовки и разделы с содержанием. "
        "Используй чёткую разметку: заголовок 1 уровня — строка, начинающаяся с «# », заголовок 2 — «## », параграфы — обычный текст."
    )


def generate_document_content(event: dict, doc_type: str | None = None) -> str:
    """
    Запрашивает у агента текст документа для события.
    event: {"date", "title", "description?", "document_type"?}
    doc_type: report | event_scenario | algorithm; если не передан, берётся из event["document_type"] или по умолчанию report.
    """
    doc_type = doc_type or event.get("document_type") or "report"
    if doc_type not in DOCUMENT_TYPES:
        doc_type = "report"
    title = event.get("title", "Событие")
    desc = event.get("description", "")
    prompt = _doc_type_prompt(doc_type, title, desc)
    return ask(prompt)


def build_docx(content: str, title: str, filename: str | None = None) -> Path:
    """
    Создаёт .docx из текста с разметкой # / ## и параграфами.
    Сохраняет в data/generated_docs/. Возвращает путь к файлу.
    """
    from docx import Document
    from docx.shared import Pt

    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, 0)

    for line in content.splitlines():
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            doc.add_paragraph(line)

    safe_name = re.sub(r'[^\w\s\-\.]', "", title)[:80].strip() or "document"
    if not filename:
        filename = f"{safe_name}.docx"
    if not filename.endswith(".docx"):
        filename += ".docx"
    out_path = GENERATED_DIR / filename
    doc.save(str(out_path))
    logger.info("Документ сохранён: %s", out_path)
    return out_path


def create_and_deliver_document(
    event: dict,
    doc_type: str | None = None,
    to_email: str | None = None,
    upload_to_vector_store: bool = True,
) -> tuple[Path | None, str]:
    """
    Генерирует документ по событию, сохраняет .docx, отправляет на почту, при возможности загружает в Vector Store.
    Возвращает (path_to_docx, сообщение об итоге).
    """
    from src.email_sender import send_email
    from src.vector_store_client import upload_file_to_vector_store

    title = event.get("title", "Документ по гранту")
    try:
        content = generate_document_content(event, doc_type)
    except Exception as e:
        logger.exception("Ошибка генерации контента: %s", e)
        return None, f"Ошибка при запросе к агенту: {e}"

    doc_type = doc_type or event.get("document_type") or "report"
    file_title = f"{title} ({doc_type})"
    try:
        path = build_docx(content, file_title)
    except Exception as e:
        logger.exception("Ошибка создания Word: %s", e)
        return None, f"Ошибка создания документа Word: {e}"

    sent = send_email(to_email=to_email, subject=f"Документ по гранту: {title}", body=f"Во вложении документ по событию: {title}.", attachments=[path])
    msg_parts = [f"Документ сохранён: {path}."]
    if sent:
        msg_parts.append("Письмо с вложением отправлено на почту.")
    else:
        msg_parts.append("Почта не отправлена (проверьте настройки SMTP и NOTIFICATION_EMAIL).")

    try:
        from src.yandex_disk_client import upload_file as disk_upload, is_configured as disk_configured
        if disk_configured() and disk_upload(path):
            msg_parts.append("Копия загружена на Яндекс.Диск.")
    except Exception as e:
        logger.debug("Яндекс.Диск: %s", e)

    if upload_to_vector_store:
        cfg = get_yandex_ai_config()
        vs_id = cfg.get("vector_store_id", "").strip()
        if vs_id:
            if upload_file_to_vector_store(path, vs_id):
                msg_parts.append("Документ добавлен в базу знаний агента (Vector Store).")
            else:
                msg_parts.append("В базу знаний агента загрузить не удалось (проверьте vector_store_id и права API).")
        else:
            msg_parts.append("vector_store_id не задан — в базу агента документ не загружен.")

    return path, " ".join(msg_parts)
